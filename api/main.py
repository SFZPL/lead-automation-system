#!/usr/bin/env python3
"""FastAPI backend exposing Perplexity and Apollo follow-up helpers."""

# Force reload
import logging
from datetime import datetime
from typing import Any, Dict, List, Optional, Literal

from fastapi import FastAPI, HTTPException, Request, Depends, UploadFile, File, Form, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, RedirectResponse, StreamingResponse, FileResponse
from pydantic import BaseModel
import json
import asyncio
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT

import os
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from config import Config
from modules.logger import setup_logging
from modules.perplexity_workflow import PerplexityWorkflow
from modules.perplexity_client import PerplexityClient
from modules.apollo_followup import ApolloFollowUpService
from modules.post_contact_automation import PostContactAutomationService, PostContactAction
from modules.lost_lead_analyzer import LostLeadAnalyzer
from modules.proposal_followup_analyzer import ProposalFollowupAnalyzer
from modules.outlook_client import OutlookClient
from modules.email_token_store import EmailTokenStore
from modules.odoo_client import OdooClient
from modules.teams_messenger import TeamsMessenger
from modules.weekly_pipeline_analyzer import WeeklyPipelineAnalyzer
from api.auth import get_auth_service, get_current_user, get_database, AuthService
from api.database import Database
from api.supabase_client import get_supabase_client
from api.supabase_database import SupabaseDatabase

logger = logging.getLogger(__name__)

# Initialize Supabase client
supabase = get_supabase_client()

# In-memory cache for proposal followups analysis (fallback if Supabase unavailable)
proposal_followups_cache = {
    "data": None,
    "timestamp": None,
    "params": None
}


def get_supabase_database() -> SupabaseDatabase:
    """Dependency to get Supabase database instance."""
    return SupabaseDatabase()


def get_user_odoo_client(user: Dict[str, Any], db: Database) -> OdooClient:
    """
    Get an Odoo client configured with the user's stored credentials.

    Args:
        user: Current authenticated user dict
        db: Database instance

    Returns:
        Configured OdooClient instance

    Raises:
        HTTPException: If user hasn't set up Odoo credentials
    """
    user_settings = db.get_user_settings(user["id"])

    if not user_settings.get("odoo_username") or not user_settings.get("odoo_password"):
        raise HTTPException(
            status_code=403,
            detail="Odoo credentials not configured. Please log in again."
        )

    # Create a custom config with user's credentials
    config = Config()
    config.ODOO_URL = user_settings.get("odoo_url") or config.ODOO_URL
    config.ODOO_DB = user_settings.get("odoo_db") or config.ODOO_DB
    config.ODOO_USERNAME = user_settings["odoo_username"]
    config.ODOO_PASSWORD = user_settings["odoo_password"]

    client = OdooClient(config)
    return client


app = FastAPI(
    title="Lead Automation API",
    description="Generate Perplexity prompts, parse results, and prepare Apollo follow-up emails.",
    version="1.1.0",
)

ALLOWED_ORIGINS = [
    'http://localhost:3000',
    'http://localhost:3001',
    'http://localhost:3002',
    'http://127.0.0.1:3000',
    'http://127.0.0.1:3001',
    'http://127.0.0.1:3002',
    'https://lead-automation-system.onrender.com',
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=['*'],
    allow_headers=['*'],
)


def _setup_logging() -> Config:
    config = Config()
    setup_logging(config, "INFO")
    return config


def get_workflow() -> PerplexityWorkflow:
    config = _setup_logging()
    return PerplexityWorkflow(config)


def get_followup_service() -> ApolloFollowUpService:
    config = _setup_logging()
    return ApolloFollowUpService(config=config)


def get_post_contact_service() -> PostContactAutomationService:
    config = _setup_logging()
    return PostContactAutomationService(config=config)


def get_lost_lead_analyzer() -> LostLeadAnalyzer:
    config = _setup_logging()
    return LostLeadAnalyzer(config=config, supabase_client=supabase)


def _to_iso(value: Any) -> Optional[str]:
    if isinstance(value, datetime):
        return value.isoformat()
    if isinstance(value, str):
        return value
    return None


def _serialize_call_info(call: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    if not call:
        return {}

    raw_call = call.get('raw_call') if isinstance(call.get('raw_call'), dict) else {}

    def _first(*keys: str) -> Optional[Any]:
        for key in keys:
            if key in call and call[key]:
                return call[key]
            if key in raw_call and raw_call[key]:
                return raw_call[key]
        return None

    result: Dict[str, Any] = {
        'id': _first('call_id', 'id'),
        'disposition': _first('call_disposition', 'disposition'),
        'duration_seconds': _first('duration_seconds', 'duration'),
        'last_called_at': _to_iso(
            _first('last_called_at_dt', 'last_called_at', 'called_at', 'updated_at')
        ),
        'notes': _first('notes', 'note'),
    }
    return {key: value for key, value in result.items() if value is not None}


def _serialize_lead_info(lead: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    if not lead:
        return {}
    result: Dict[str, Any] = {
        'id': lead.get('id'),
        'name': lead.get('name') or lead.get('contact_name'),
        'company': lead.get('partner_name') or lead.get('Company Name') or lead.get('company'),
        'stage_name': lead.get('stage_name'),
        'salesperson': lead.get('salesperson_name') or lead.get('Salesperson'),
        'phone': lead.get('phone') or lead.get('mobile'),
    }
    return {key: value for key, value in result.items() if value is not None}


class LeadPreview(BaseModel):
    id: int
    full_name: Optional[str] = None
    company_name: Optional[str] = None
    email: Optional[str] = None
    linkedin: Optional[str] = None


class GenerateResponse(BaseModel):
    prompt: str
    lead_count: int
    leads: List[LeadPreview]


class ParseRequest(BaseModel):
    results_text: str
    update: bool = True


class ParseResponse(BaseModel):
    parsed_count: int
    updated: int
    failed: int
    errors: List[str]


class EnrichSingleLeadRequest(BaseModel):
    lead_id: int


class EnrichSingleLeadResponse(BaseModel):
    success: bool
    lead_id: int
    enriched_data: Optional[Dict[str, Any]] = None
    error: Optional[str] = None


class EnrichBatchRequest(BaseModel):
    lead_ids: List[int]


class EnrichedLeadResult(BaseModel):
    lead_id: int
    success: bool
    current_data: Optional[Dict[str, Any]] = None
    suggested_data: Optional[Dict[str, Any]] = None
    error: Optional[str] = None


class ParsePreviewRequest(BaseModel):
    results_text: str


class ParsePreviewResponse(BaseModel):
    success: bool
    results: List[EnrichedLeadResult]


class LeadComplexityAnalysis(BaseModel):
    lead_id: int
    name: str
    complexity_score: int
    complexity_level: str
    factors: List[str]


class BatchRecommendation(BaseModel):
    batch_number: int
    lead_count: int
    total_complexity: int
    leads: List[Dict[str, Any]]
    prompt: str


class SmartAnalysisResponse(BaseModel):
    total_leads: int
    recommended_batches: int
    batches: List[BatchRecommendation]
    analysis_summary: str  # Reuse the same model


class EnrichBatchResponse(BaseModel):
    total: int
    successful: int
    failed: int
    results: List[EnrichedLeadResult]


class PushApprovedRequest(BaseModel):
    approved_leads: List[Dict[str, Any]]
    send_emails: Optional[bool] = False
    email_data: Optional[Dict[int, Dict[str, str]]] = None  # lead_id -> {subject, body}


class PushApprovedResponse(BaseModel):
    total: int
    successful: int
    failed: int
    errors: List[str]
    emails_sent: Optional[int] = 0
    email_errors: Optional[List[str]] = []


class FollowUpCall(BaseModel):
    id: Optional[str] = None
    disposition: Optional[str] = None
    direction: Optional[str] = None
    duration_seconds: Optional[int] = None
    last_called_at: Optional[str] = None
    notes: Optional[str] = None


class FollowUpLead(BaseModel):
    id: Optional[int] = None
    name: Optional[str] = None
    stage_name: Optional[str] = None
    salesperson: Optional[str] = None
    phone: Optional[str] = None
    company: Optional[str] = None


class FollowUpItem(BaseModel):
    email: str
    subject: str
    body: str
    call: FollowUpCall
    odoo_lead: FollowUpLead


class FollowUpResponse(BaseModel):
    count: int
    items: List[FollowUpItem]


class PostContactCall(BaseModel):
    id: Optional[str] = None
    disposition: Optional[str] = None
    duration_seconds: Optional[int] = None
    last_called_at: Optional[str] = None
    notes: Optional[str] = None


class PostContactLead(BaseModel):
    id: Optional[int] = None
    name: Optional[str] = None
    company: Optional[str] = None
    stage_name: Optional[str] = None
    salesperson: Optional[str] = None
    phone: Optional[str] = None


class PostContactActionPayload(BaseModel):
    action_type: Literal["email", "note"]
    contact_email: str
    odoo_lead_id: int
    subject: Optional[str] = None
    body: Optional[str] = None
    note_body: Optional[str] = None
    transcription: Optional[str] = None
    call: Optional[PostContactCall] = None
    odoo_lead: Optional[PostContactLead] = None


class PostContactActionsResponse(BaseModel):
    count: int
    actions: List[PostContactActionPayload]


class ExecuteActionRequest(PostContactActionPayload):
    pass


class ExecuteActionResponse(BaseModel):
    success: bool
    message: Optional[str] = None


class LostLeadSummary(BaseModel):
    id: int
    name: str
    record_type: Optional[str] = None  # 'lead' or 'opportunity'
    partner_name: Optional[str] = None
    contact_name: Optional[str] = None
    stage: Optional[str] = None
    lost_reason: Optional[str] = None
    lost_reason_category: Optional[str] = None
    probability: Optional[float] = None
    expected_revenue: Optional[float] = None
    salesperson: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    mobile: Optional[str] = None
    create_date: Optional[str] = None
    last_update: Optional[str] = None


class LostLeadListResponse(BaseModel):
    count: int
    items: List[LostLeadSummary]

    class Config:
        # Ensure None values are included in JSON response
        json_encoders = {}
        use_enum_values = True


class LostLeadAnalysisRequest(BaseModel):
    max_internal_notes: Optional[int] = None
    max_emails: Optional[int] = None
    user_identifier: Optional[str] = None  # Email address for Outlook search
    include_outlook_emails: bool = False  # Whether to search Outlook emails


class LostLeadMessage(BaseModel):
    id: Optional[int]
    date: Optional[str]
    formatted_date: Optional[str]
    author: Optional[str]
    subject: Optional[str]
    body: str


class LostLeadAnalysisResponse(BaseModel):
    lead: Dict[str, Any]
    analysis: Dict[str, Any]
    internal_notes: List[LostLeadMessage]
    emails: List[LostLeadMessage]


@app.get("/health")
def health() -> dict:
    return {"status": "ok"}


@app.get("/test-type")
def test_type() -> dict:
    """Test endpoint to debug type field issue."""
    from pydantic import BaseModel
    from typing import Optional

    class TestModel(BaseModel):
        id: int
        record_type: Optional[str] = None
        type: Optional[str] = None

    obj = TestModel(id=1, record_type="test_record", type="test_type")
    return {
        "object": obj.dict(),
        "pydantic_version": __import__("pydantic").VERSION
    }


@app.post("/perplexity/generate", response_model=GenerateResponse)
def generate_prompt() -> GenerateResponse:
    workflow = get_workflow()

    prompt, leads = workflow.generate_enrichment_prompt()

    if not leads:
        return GenerateResponse(prompt="", lead_count=0, leads=[])

    previews = [
        LeadPreview(
            id=lead.get("id"),
            full_name=lead.get("Full Name") or lead.get("name"),
            company_name=lead.get("Company Name") or lead.get("partner_name"),
            email=lead.get("email"),
            linkedin=lead.get("LinkedIn Link"),
        )
        for lead in leads
        if lead.get("id") is not None
    ]

    return GenerateResponse(prompt=prompt, lead_count=len(leads), leads=previews)


@app.post("/perplexity/smart-analysis", response_model=SmartAnalysisResponse)
def smart_analysis() -> SmartAnalysisResponse:
    """Analyze leads and recommend optimal batch splitting for manual enrichment"""
    workflow = get_workflow()

    # Get all unenriched leads
    _, leads = workflow.generate_enrichment_prompt()
    if not leads:
        return SmartAnalysisResponse(
            total_leads=0,
            recommended_batches=0,
            batches=[],
            analysis_summary="No unenriched leads found."
        )

    # Get optimized batch split
    batches = workflow.optimize_batch_split(leads)

    # Build response with prompts for each batch
    batch_recommendations = []
    for batch in batches:
        # Extract just the leads (not the analysis wrapper)
        batch_leads = [item['lead'] for item in batch['leads']]

        # Generate prompt for this batch
        batch_prompt = workflow._build_comprehensive_prompt(batch_leads)

        # Build lead info with complexity
        leads_info = []
        for item in batch['leads']:
            leads_info.append({
                "id": item['lead'].get('id'),
                "name": item['analysis']['name'],
                "complexity": item['analysis']['complexity_level'],
                "factors": item['analysis']['factors']
            })

        batch_recommendations.append(BatchRecommendation(
            batch_number=batch['batch_number'],
            lead_count=batch['lead_count'],
            total_complexity=batch['total_complexity'],
            leads=leads_info,
            prompt=batch_prompt
        ))

    # Generate summary
    summary_parts = [
        f"📊 **Smart Analysis Complete**",
        f"",
        f"Total Leads: {len(leads)}",
        f"Recommended Batches: {len(batches)}",
        f"",
        f"**Strategy:** Leads are sorted by complexity and grouped to balance search depth with batch size.",
        f"High-complexity leads (missing info, common names) are prioritized for focused attention.",
        f"",
        f"**Why Split?** Smaller batches = deeper research per lead, more accurate results, and better LinkedIn URL verification."
    ]

    summary = "\n".join(summary_parts)

    return SmartAnalysisResponse(
        total_leads=len(leads),
        recommended_batches=len(batches),
        batches=batch_recommendations,
        analysis_summary=summary
    )


@app.post("/perplexity/parse", response_model=ParseResponse)
def parse_results(payload: ParseRequest) -> ParseResponse:
    workflow = get_workflow()

    _, original_leads = workflow.generate_enrichment_prompt()
    if not original_leads:
        raise HTTPException(status_code=409, detail="No leads available to reconcile the results against")

    enriched = workflow.parse_perplexity_results(payload.results_text, original_leads)
    if not enriched:
        raise HTTPException(status_code=422, detail="Unable to parse Perplexity response")

    updated = 0
    failed = 0
    errors: List[str] = []

    if payload.update:
        outcome = workflow.update_leads_in_odoo(enriched)
        if not outcome.get("success", False):
            raise HTTPException(status_code=500, detail=outcome.get("error", "Unknown error"))
        updated = outcome.get("updated", 0)
        failed = outcome.get("failed", 0)
        errors = outcome.get("errors", [])

    return ParseResponse(
        parsed_count=len(enriched),
        updated=updated,
        failed=failed,
        errors=errors,
    )


@app.post("/perplexity/parse-preview", response_model=ParsePreviewResponse)
def parse_preview(payload: ParsePreviewRequest) -> ParsePreviewResponse:
    """Parse Perplexity output and return preview without pushing to Odoo"""
    workflow = get_workflow()

    _, original_leads = workflow.generate_enrichment_prompt()
    if not original_leads:
        raise HTTPException(status_code=409, detail="No leads available to reconcile the results against")

    enriched = workflow.parse_perplexity_results(payload.results_text, original_leads)
    if not enriched:
        raise HTTPException(status_code=422, detail="Unable to parse Perplexity response")

    # Build results for preview
    results = []
    for enriched_lead in enriched:
        # Find original lead to show current vs suggested
        original = next((l for l in original_leads if l.get('id') == enriched_lead.get('id')), {})

        results.append(EnrichedLeadResult(
            lead_id=enriched_lead.get('id', 0),
            success=True,
            current_data=original,
            suggested_data=enriched_lead
        ))

    return ParsePreviewResponse(
        success=True,
        results=results
    )


@app.post("/perplexity/enrich-single", response_model=EnrichSingleLeadResponse)
def enrich_single_lead(payload: EnrichSingleLeadRequest) -> EnrichSingleLeadResponse:
    """Enrich a single lead using Perplexity API"""
    try:
        config = Config()
        workflow = PerplexityWorkflow(config)
        perplexity_client = PerplexityClient(config)

        # Connect to Odoo and get the lead
        if not workflow.odoo.connect():
            return EnrichSingleLeadResponse(
                success=False,
                lead_id=payload.lead_id,
                error="Failed to connect to Odoo"
            )

        # Fetch the specific lead by ID
        lead_data = workflow.odoo._call_kw(
            'crm.lead', 'read',
            [[payload.lead_id]],
            {'fields': [
                'id', 'name', 'partner_name', 'email_from', 'phone', 'mobile',
                'function', 'contact_name', 'x_studio_linkedin_profile'
            ]}
        )

        if not lead_data:
            return EnrichSingleLeadResponse(
                success=False,
                lead_id=payload.lead_id,
                error=f"Lead {payload.lead_id} not found"
            )

        # Convert to the format expected by the workflow
        lead = lead_data[0]
        formatted_lead = {
            'id': lead.get('id'),
            'Full Name': lead.get('name') or lead.get('contact_name') or '',
            'Company Name': lead.get('partner_name') or '',
            'email': lead.get('email_from') or '',
            'Phone': lead.get('phone') or '',
            'Mobile': lead.get('mobile') or '',
            'Job Role': lead.get('function') or '',
        }

        # Generate prompt for this single lead
        prompt = workflow.generate_single_lead_prompt(formatted_lead)
        logger.info(f"Generated enrichment prompt for lead {payload.lead_id}")

        # Call Perplexity API
        perplexity_response = perplexity_client.enrich_lead(formatted_lead, prompt)

        if not perplexity_response:
            return EnrichSingleLeadResponse(
                success=False,
                lead_id=payload.lead_id,
                error="Perplexity API returned no response"
            )

        # Parse the response
        enriched_leads = workflow.parse_perplexity_results(perplexity_response, [formatted_lead])

        if not enriched_leads:
            return EnrichSingleLeadResponse(
                success=False,
                lead_id=payload.lead_id,
                error="Failed to parse Perplexity response"
            )

        enriched_data = enriched_leads[0]
        logger.info(f"Successfully enriched lead {payload.lead_id}")

        return EnrichSingleLeadResponse(
            success=True,
            lead_id=payload.lead_id,
            enriched_data=enriched_data
        )

    except Exception as e:
        logger.error(f"Error enriching lead {payload.lead_id}: {str(e)}")
        return EnrichSingleLeadResponse(
            success=False,
            lead_id=payload.lead_id,
            error=str(e)
        )


@app.post("/perplexity/enrich-batch", response_model=EnrichBatchResponse)
def enrich_batch_leads(payload: EnrichBatchRequest) -> EnrichBatchResponse:
    """Enrich multiple leads using Perplexity API in ONE batch call - returns comparison data for review"""
    results = []
    successful = 0
    failed = 0

    config = Config()
    workflow = PerplexityWorkflow(config)
    perplexity_client = PerplexityClient(config)

    # Connect to Odoo
    if not workflow.odoo.connect():
        return EnrichBatchResponse(
            total=len(payload.lead_ids),
            successful=0,
            failed=len(payload.lead_ids),
            results=[
                EnrichedLeadResult(
                    lead_id=lead_id,
                    success=False,
                    error="Failed to connect to Odoo"
                )
                for lead_id in payload.lead_ids
            ]
        )

    try:
        # Fetch ALL leads from Odoo
        all_leads_data = workflow.odoo._call_kw(
            'crm.lead', 'read',
            [payload.lead_ids],
            {'fields': [
                'id', 'name', 'partner_name', 'email_from', 'phone', 'mobile',
                'function', 'contact_name', 'x_studio_linkedin_profile',
                'website', 'city', 'country_id', 'x_studio_quality'
            ]}
        )

        # Build map of current data and formatted leads
        current_data_map = {}
        formatted_leads = []

        for lead in all_leads_data:
            lead_id = lead.get('id')

            # Store current data
            current_data_map[lead_id] = {
                'id': lead_id,
                'Full Name': lead.get('name') or lead.get('contact_name') or '',
                'Company Name': lead.get('partner_name') or '',
                'email': lead.get('email_from') or '',
                'Phone': lead.get('phone') or '',
                'Mobile': lead.get('mobile') or '',
                'Job Role': lead.get('function') or '',
                'LinkedIn Link': lead.get('x_studio_linkedin_profile') or '',
                'website': lead.get('website') or '',
                'City': lead.get('city') or '',
                'Country': lead.get('country_id')[1] if lead.get('country_id') else '',
                'Quality (Out of 5)': lead.get('x_studio_quality') or '',
            }

            # Format for enrichment
            formatted_leads.append({
                'id': lead_id,
                'Full Name': lead.get('name') or lead.get('contact_name') or '',
                'Company Name': lead.get('partner_name') or '',
                'email': lead.get('email_from') or '',
                'Phone': lead.get('phone') or '',
                'Mobile': lead.get('mobile') or '',
                'Job Role': lead.get('function') or '',
            })

        # Generate ONE batch prompt for ALL leads
        logger.info(f"Generating batch prompt for {len(formatted_leads)} leads")
        batch_prompt = workflow._build_comprehensive_prompt(formatted_leads)

        # Call Perplexity ONCE with all leads
        logger.info(f"Calling Perplexity API with batch of {len(formatted_leads)} leads")
        perplexity_response = perplexity_client.search(batch_prompt, max_tokens=8000)

        # DEBUG: Save full raw response
        if perplexity_response:
            with open('c:/Users/Geeks/Desktop/Programming_Files/Leads/perplexity_raw_response.txt', 'w', encoding='utf-8') as f:
                f.write(perplexity_response)
            logger.info(f"Perplexity raw response saved to perplexity_raw_response.txt ({len(perplexity_response)} chars)")

        if not perplexity_response:
            # All leads failed
            for lead_id in payload.lead_ids:
                results.append(EnrichedLeadResult(
                    lead_id=lead_id,
                    success=False,
                    current_data=current_data_map.get(lead_id),
                    error="Perplexity API returned no response"
                ))
            return EnrichBatchResponse(
                total=len(payload.lead_ids),
                successful=0,
                failed=len(payload.lead_ids),
                results=results
            )

        # Parse batch response (this works reliably!)
        logger.info(f"Parsing batch response ({len(perplexity_response)} characters)")
        enriched_leads = workflow.parse_perplexity_results(perplexity_response, formatted_leads)

        # Match enriched leads back to lead_ids
        enriched_map = {lead.get('id'): lead for lead in enriched_leads}

        for lead_id in payload.lead_ids:
            if lead_id in enriched_map:
                results.append(EnrichedLeadResult(
                    lead_id=lead_id,
                    success=True,
                    current_data=current_data_map.get(lead_id),
                    suggested_data=enriched_map[lead_id]
                ))
                successful += 1
            else:
                results.append(EnrichedLeadResult(
                    lead_id=lead_id,
                    success=False,
                    current_data=current_data_map.get(lead_id),
                    error="Lead not found in Perplexity response"
                ))
                failed += 1

        logger.info(f"Batch enrichment complete: {successful} successful, {failed} failed")

    except Exception as e:
        logger.error(f"Error in batch enrichment: {str(e)}")
        # Return error for all leads
        for lead_id in payload.lead_ids:
            results.append(EnrichedLeadResult(
                lead_id=lead_id,
                success=False,
                error=str(e)
            ))
        failed = len(payload.lead_ids)

    return EnrichBatchResponse(
        total=len(payload.lead_ids),
        successful=successful,
        failed=failed,
        results=results
    )


@app.post("/perplexity/enrich-batch-stream")
async def enrich_batch_stream(payload: EnrichBatchRequest):
    """Enrich leads individually with streaming progress updates"""

    async def event_generator():
        config = Config()
        workflow = PerplexityWorkflow(config)
        perplexity_client = PerplexityClient(config)

        # Connect to Odoo
        if not workflow.odoo.connect():
            yield f"data: {json.dumps({'type': 'error', 'message': 'Failed to connect to Odoo'})}\n\n"
            return

        results = []
        successful = 0
        failed = 0

        try:
            # Fetch ALL leads from Odoo first
            all_leads_data = workflow.odoo._call_kw(
                'crm.lead', 'read',
                [payload.lead_ids],
                {'fields': [
                    'id', 'name', 'partner_name', 'email_from', 'phone', 'mobile',
                    'function', 'contact_name', 'x_studio_linkedin_profile',
                    'website', 'city', 'country_id', 'x_studio_quality'
                ]}
            )

            # Create a map for easy lookup
            leads_map = {lead['id']: lead for lead in all_leads_data}

            # Enrich each lead individually
            for i, lead_id in enumerate(payload.lead_ids):
                lead = leads_map.get(lead_id)
                if not lead:
                    yield f"data: {json.dumps({'type': 'error', 'lead_id': lead_id, 'message': f'Lead {lead_id} not found'})}\n\n"
                    failed += 1
                    continue

                lead_name = lead.get('name') or lead.get('contact_name') or f'Lead {lead_id}'

                # Send progress update
                yield f"data: {json.dumps({'type': 'progress', 'lead_id': lead_id, 'lead_name': lead_name, 'current': i + 1, 'total': len(payload.lead_ids)})}\n\n"

                # Format lead for enrichment
                formatted_lead = {
                    'id': lead_id,
                    'Full Name': lead.get('name') or lead.get('contact_name') or '',
                    'Company Name': lead.get('partner_name') or '',
                    'email': lead.get('email_from') or '',
                    'Phone': lead.get('phone') or '',
                    'Mobile': lead.get('mobile') or '',
                    'Job Role': lead.get('function') or '',
                }

                try:
                    # Generate prompt for single lead
                    prompt = workflow.generate_single_lead_prompt(formatted_lead)

                    # Call Perplexity
                    perplexity_response = perplexity_client.search(prompt)

                    # Parse response
                    enriched_data = workflow.parse_single_lead_response(perplexity_response, formatted_lead)

                    # Store current data
                    current_data = {
                        'id': lead_id,
                        'Full Name': lead.get('name') or lead.get('contact_name') or '',
                        'Company Name': lead.get('partner_name') or '',
                        'email': lead.get('email_from') or '',
                        'Phone': lead.get('phone') or '',
                        'Mobile': lead.get('mobile') or '',
                        'Job Role': lead.get('function') or '',
                        'LinkedIn Link': lead.get('x_studio_linkedin_profile') or '',
                        'website': lead.get('website') or '',
                        'City': lead.get('city') or '',
                        'Country': lead.get('country_id')[1] if lead.get('country_id') else '',
                        'Quality (Out of 5)': lead.get('x_studio_quality') or '',
                    }

                    result = EnrichedLeadResult(
                        lead_id=lead_id,
                        success=True,
                        current_data=current_data,
                        suggested_data=enriched_data
                    )
                    results.append(result)
                    successful += 1

                    # Send success update
                    yield f"data: {json.dumps({'type': 'success', 'lead_id': lead_id, 'lead_name': lead_name})}\n\n"

                except Exception as e:
                    logger.error(f"Error enriching lead {lead_id}: {e}")
                    result = EnrichedLeadResult(
                        lead_id=lead_id,
                        success=False,
                        error=str(e)
                    )
                    results.append(result)
                    failed += 1

                    # Send error update
                    yield f"data: {json.dumps({'type': 'error', 'lead_id': lead_id, 'lead_name': lead_name, 'message': str(e)})}\n\n"

                # Small delay to avoid rate limiting
                await asyncio.sleep(1)

            # Send final completion
            final_response = {
                'type': 'complete',
                'total': len(payload.lead_ids),
                'successful': successful,
                'failed': failed,
                'results': [r.dict() for r in results]
            }
            yield f"data: {json.dumps(final_response)}\n\n"

        except Exception as e:
            logger.error(f"Error in batch enrichment stream: {e}")
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")


@app.post("/perplexity/push-approved", response_model=PushApprovedResponse)
def push_approved_enrichments(
    payload: PushApprovedRequest,
    current_user: Dict[str, Any] = Depends(get_current_user)
) -> PushApprovedResponse:
    """Push approved enrichment data to Odoo and optionally send outreach emails"""
    config = Config()
    workflow = PerplexityWorkflow(config)

    if not payload.approved_leads:
        return PushApprovedResponse(
            total=0,
            successful=0,
            failed=0,
            errors=["No leads provided"]
        )

    logger.info(f"Pushing {len(payload.approved_leads)} approved leads to Odoo")

    try:
        outcome = workflow.update_leads_in_odoo(payload.approved_leads)
        print("[EMAIL-DEBUG-0-PRINT] Odoo update completed - checking outcome")
        logger.info(f"[EMAIL-DEBUG-0] Odoo update returned. success={outcome.get('success')}, type={type(outcome)}")

        if not outcome.get("success", False):
            return PushApprovedResponse(
                total=len(payload.approved_leads),
                successful=0,
                failed=len(payload.approved_leads),
                errors=[outcome.get("error", "Unknown error")]
            )

        updated = outcome.get("updated", 0)
        failed = outcome.get("failed", 0)
        errors = outcome.get("errors", [])

        logger.info(f"[EMAIL-DEBUG-1] Successfully pushed {updated} leads to Odoo, {failed} failed")
        logger.info(f"[EMAIL-DEBUG-2] About to check email sending. send_emails={payload.send_emails}, has_email_data={bool(payload.email_data)}")

        # Send emails if requested
        emails_sent = 0
        email_errors = []

        logger.info(f"[EMAIL-DEBUG-3] Email sending request: send_emails={payload.send_emails}, email_data_count={len(payload.email_data) if payload.email_data else 0}")

        if payload.send_emails and payload.email_data:
            from modules.outlook_client import OutlookClient
            import os

            outlook = OutlookClient()
            pdf_path = os.path.join(os.path.dirname(__file__), '..', 'knowledge_base', 'PrezLab Company Profile.pdf')

            # Get authenticated user's Outlook access token
            # Use current user's ID as identifier (matches the auth flow)
            user_identifier = str(current_user["id"])

            try:
                logger.info(f"🔐 Retrieving Outlook tokens for user {user_identifier}")
                tokens = outlook.get_user_auth_tokens(user_identifier)
                if not tokens or 'access_token' not in tokens:
                    logger.error(f"❌ No Outlook tokens found for user {user_identifier}")
                    email_errors.append(f"No authenticated Outlook account found for {current_user['email']}. Please authenticate your Outlook account in Settings.")
                else:
                    access_token = tokens['access_token']
                    logger.info(f"✅ Got Outlook access token for user {user_identifier}, preparing to send {len(payload.approved_leads)} emails")
                    logger.info(f"📋 Email data keys: {list(payload.email_data.keys())}")
                    logger.info(f"📋 Lead IDs to process: {[lead.get('id') for lead in payload.approved_leads]}")

                    for lead in payload.approved_leads:
                        lead_id = lead.get('id')
                        lead_email = lead.get('email')

                        if not lead_id or not lead_email:
                            logger.warning(f"⚠️ Skipping lead: missing ID or email (id={lead_id}, email={lead_email})")
                            continue

                        # Try both integer and string keys (frontend sends as int)
                        email_draft = payload.email_data.get(lead_id) or payload.email_data.get(str(lead_id))
                        if not email_draft:
                            logger.warning(f"⚠️ No email draft found for lead {lead_id} (tried both int and str keys)")
                            continue

                        subject = email_draft.get('subject', 'Thank you for your interest in PrezLab')
                        body = email_draft.get('body', '').replace('\n', '<br>')

                        logger.info(f"📤 Sending email to {lead_email} (Lead {lead_id}): '{subject}'")

                        try:
                            success = outlook.send_email_with_attachment(
                                access_token=access_token,
                                to=[lead_email],
                                subject=subject,
                                body=body,
                                attachment_path=pdf_path,
                                attachment_name='PrezLab Company Profile.pdf',
                                cc=['engage@prezlab.com']
                            )

                            if success:
                                emails_sent += 1
                                logger.info(f"Email sent successfully to {lead_email} from {current_user['email']}")
                            else:
                                email_errors.append(f"Failed to send email to {lead_email}")

                        except Exception as email_error:
                            logger.error(f"Error sending email to {lead_email}: {email_error}")
                            email_errors.append(f"Error sending to {lead_email}: {str(email_error)}")

            except Exception as auth_error:
                logger.error(f"Error getting Outlook authentication: {auth_error}")
                email_errors.append(f"Authentication error: {str(auth_error)}")

        return PushApprovedResponse(
            total=len(payload.approved_leads),
            successful=updated,
            failed=failed,
            errors=errors,
            emails_sent=emails_sent,
            email_errors=email_errors
        )

    except Exception as e:
        logger.error(f"Error pushing approved leads to Odoo: {str(e)}")
        return PushApprovedResponse(
            total=len(payload.approved_leads),
            successful=0,
            failed=len(payload.approved_leads),
            errors=[str(e)]
        )


@app.get("/apollo/followups", response_model=FollowUpResponse)
def get_apollo_followups(limit: int = 10, lookback_hours: Optional[int] = None) -> FollowUpResponse:
    service = get_followup_service()
    try:
        followups = service.prepare_followups(limit=limit, lookback_hours=lookback_hours)
    except Exception as exc:
        logger.error("Failed to prepare Apollo follow-ups: %s", exc)
        raise HTTPException(status_code=500, detail="Unable to generate Apollo follow-ups")

    items: List[FollowUpItem] = []
    for item in followups:
        email = item.get('email')
        subject = item.get('subject', '')
        body = item.get('body', '')
        if not email or not subject or not body:
            continue

        call = item.get('call') or {}
        odoo_lead = item.get('odoo_lead') or {}
        items.append(
            FollowUpItem(
                email=email,
                subject=subject,
                body=body,
                call=FollowUpCall(
                    id=call.get('id'),
                    disposition=call.get('disposition'),
                    direction=call.get('direction'),
                    duration_seconds=call.get('duration_seconds'),
                    last_called_at=call.get('last_called_at'),
                    notes=call.get('notes'),
                ),
                odoo_lead=FollowUpLead(
                    id=odoo_lead.get('id'),
                    name=odoo_lead.get('name'),
                    stage_name=odoo_lead.get('stage_name'),
                    salesperson=odoo_lead.get('salesperson'),
                    phone=str(odoo_lead.get('phone')) if odoo_lead.get('phone') else None,
                    company=odoo_lead.get('company'),
                ),
            )
        )

    return FollowUpResponse(count=len(items), items=items)


@app.get("/post-contact/actions", response_model=PostContactActionsResponse)
def get_post_contact_actions(limit: int = 10, lookback_hours: Optional[int] = None) -> PostContactActionsResponse:
    service = get_post_contact_service()
    try:
        actions = service.prepare_actions(limit=limit, lookback_hours=lookback_hours)
    except Exception as exc:
        logger.error("Failed to prepare post-contact actions: %s", exc)
        raise HTTPException(status_code=500, detail="Unable to prepare post-contact actions")

    serialized: List[PostContactActionPayload] = []
    for action in actions:
        call_info = _serialize_call_info(action.call)
        lead_info = _serialize_lead_info(action.odoo_lead)

        serialized.append(
            PostContactActionPayload(
                action_type=action.action_type,  # type: ignore[arg-type]
                contact_email=action.contact_email,
                odoo_lead_id=int(action.odoo_lead_id),
                subject=action.subject,
                body=action.body,
                note_body=action.note_body,
                transcription=action.transcription,
                call=PostContactCall(**call_info) if call_info else None,
                odoo_lead=PostContactLead(**lead_info) if lead_info else None,
            )
        )

    return PostContactActionsResponse(count=len(serialized), actions=serialized)


@app.post("/post-contact/execute", response_model=ExecuteActionResponse)
def execute_post_contact_action(payload: ExecuteActionRequest) -> ExecuteActionResponse:
    service = get_post_contact_service()
    action = PostContactAction(
        action_type=payload.action_type,
        contact_email=payload.contact_email,
        odoo_lead_id=payload.odoo_lead_id,
        subject=payload.subject,
        body=payload.body,
        note_body=payload.note_body,
        transcription=payload.transcription,
        call=payload.call.dict() if payload.call else {},
        odoo_lead=payload.odoo_lead.dict() if payload.odoo_lead else {},
    )

    try:
        if payload.action_type == "email":
            success = service.execute_email(action)
            message = "Email sent successfully"
        else:
            success = service.execute_note(action)
            message = "Note uploaded to Odoo"
    except Exception as exc:
        logger.error("Failed to execute post-contact action: %s", exc)
        raise HTTPException(status_code=500, detail="Failed to execute post-contact action")

    if not success:
        raise HTTPException(status_code=500, detail="Action did not complete successfully")

    return ExecuteActionResponse(success=True, message=message)


@app.get("/lost-leads", response_model=LostLeadListResponse, response_model_exclude_none=False)
def list_lost_leads(
    limit: int = 20,
    salesperson: Optional[str] = None,
    type_filter: Optional[str] = None
) -> LostLeadListResponse:
    analyzer = get_lost_lead_analyzer()
    try:
        leads = analyzer.list_lost_leads(
            limit=limit,
            salesperson_name=salesperson,
            type_filter=type_filter
        )
    except Exception as exc:
        logger.error("Failed to fetch lost leads: %s", exc)
        raise HTTPException(status_code=500, detail="Unable to fetch lost leads")

    items: List[LostLeadSummary] = []
    for lead in leads:
        try:
            lead_id = int(lead.get("id"))
        except Exception:
            logger.debug("Skipping lead with missing id: %s", lead)
            continue
        def _str_or_none(value):
            """Convert Odoo False values to None for optional string fields."""
            return value if isinstance(value, str) else None

        lead_type = _str_or_none(lead.get("type"))
        logger.debug(f"Lead {lead_id}: raw type={lead.get('type')}, processed type={lead_type}")

        items.append(
            LostLeadSummary(
                id=lead_id,
                name=lead.get("name") or "Untitled Opportunity",
                record_type=lead_type,
                partner_name=_str_or_none(lead.get("partner_name")),
                contact_name=_str_or_none(lead.get("contact_name")),
                stage=_str_or_none(lead.get("stage_id")),
                lost_reason=_str_or_none(lead.get("lost_reason")),
                lost_reason_category=_str_or_none(lead.get("lost_reason_id")),
                probability=lead.get("probability"),
                expected_revenue=lead.get("expected_revenue"),
                salesperson=_str_or_none(lead.get("user_id")),
                email=_str_or_none(lead.get("email_from")),
                phone=_str_or_none(lead.get("phone")),
                mobile=_str_or_none(lead.get("mobile")),
                create_date=_str_or_none(lead.get("create_date")),
                last_update=_str_or_none(lead.get("write_date")),
            )
        )

    response = LostLeadListResponse(count=len(items), items=items)
    if items:
        logger.info(f"First item record_type field: {items[0].record_type}")
        logger.info(f"First item dict: {items[0].dict()}")
    return response


@app.post("/lost-leads/{lead_id}/analysis", response_model=LostLeadAnalysisResponse)
def analyze_lost_lead(
    lead_id: int,
    payload: LostLeadAnalysisRequest,
) -> LostLeadAnalysisResponse:
    analyzer = get_lost_lead_analyzer()
    try:
        result = analyzer.analyze_lost_lead(
            lead_id,
            max_internal_notes=payload.max_internal_notes,
            max_emails=payload.max_emails,
            user_identifier=payload.user_identifier,
            include_outlook_emails=payload.include_outlook_emails,
        )
    except RuntimeError as exc:
        raise HTTPException(status_code=400, detail=str(exc))
    except Exception as exc:
        logger.error("Failed to analyze lost lead %s: %s", lead_id, exc)
        raise HTTPException(status_code=500, detail="Unable to analyze lost lead")

    def _map_messages(messages: List[Dict[str, Any]]) -> List[LostLeadMessage]:
        mapped: List[LostLeadMessage] = []
        for message in messages:
            # Ensure all fields are proper types
            subject = message.get("subject")
            if subject is None or subject is False:
                subject = ""
            elif not isinstance(subject, str):
                subject = str(subject)

            mapped.append(
                LostLeadMessage(
                    id=message.get("id"),
                    date=message.get("date"),
                    formatted_date=message.get("formatted_date"),
                    author=message.get("author") or "",
                    subject=subject,
                    body=message.get("body") or "",
                )
            )
        return mapped

    return LostLeadAnalysisResponse(
        lead=result.get("lead", {}),
        analysis=result.get("analysis", {}),
        internal_notes=_map_messages(result.get("internal_notes", [])),
        emails=_map_messages(result.get("emails", [])),
    )


class SaveAnalysisRequest(BaseModel):
    """Request to save an analysis for re-engagement."""
    lead_id: int
    analysis_data: Dict[str, Any]
    title: Optional[str] = None


class SavedAnalysisResponse(BaseModel):
    """Response with saved analysis details."""
    success: bool
    analysis_id: Optional[str] = None
    message: Optional[str] = None


class SharedAnalysisItem(BaseModel):
    """A shared analysis item for re-engagement."""
    id: str
    lead_id: int
    title: str
    analysis_data: Dict[str, Any]
    created_by_user_id: int
    created_at: str
    lead_name: Optional[str] = None
    company_name: Optional[str] = None


@app.post("/lost-leads/{lead_id}/save-analysis", response_model=SavedAnalysisResponse)
def save_lost_lead_analysis(
    lead_id: int,
    payload: SaveAnalysisRequest,
    current_user: Dict[str, Any] = Depends(get_current_user)
) -> SavedAnalysisResponse:
    """Save a lost lead analysis for re-engagement (visible to all users)."""
    if not supabase:
        raise HTTPException(status_code=503, detail="Database not available")

    try:
        # Prepare analysis data with lead info
        analysis_to_save = {
            **payload.analysis_data,
            "lead_id": lead_id,
            "title": payload.title or f"Lost Lead #{lead_id} Analysis"
        }

        # Insert into analysis_cache with is_shared=true
        result = supabase.client.table("analysis_cache").insert({
            "user_id": current_user["id"],
            "analysis_type": "lost_leads",
            "parameters": {"lead_id": lead_id},
            "results": analysis_to_save,
            "is_shared": True  # Make visible to all users
        }).execute()

        if result.data:
            return SavedAnalysisResponse(
                success=True,
                analysis_id=result.data[0]["id"],
                message="Analysis saved and shared successfully"
            )
        else:
            raise HTTPException(status_code=500, detail="Failed to save analysis")

    except Exception as exc:
        logger.error(f"Error saving analysis for lead {lead_id}: {exc}")
        raise HTTPException(status_code=500, detail=str(exc))


@app.get("/re-engage/analyses")
def get_shared_analyses(
    current_user: Dict[str, Any] = Depends(get_current_user)
) -> List[SharedAnalysisItem]:
    """Get all shared lost lead analyses for re-engagement."""
    if not supabase:
        raise HTTPException(status_code=503, detail="Database not available")

    try:
        # Fetch all shared lost_leads analyses
        result = supabase.client.table("analysis_cache")\
            .select("*")\
            .eq("analysis_type", "lost_leads")\
            .eq("is_shared", True)\
            .order("created_at", desc=True)\
            .execute()

        if not result.data:
            return []

        # Map to response format
        shared_analyses = []
        for item in result.data:
            results_data = item.get("results", {})
            shared_analyses.append(SharedAnalysisItem(
                id=item["id"],
                lead_id=results_data.get("lead_id", 0),
                title=results_data.get("title", f"Analysis #{item['id'][:8]}"),
                analysis_data=results_data,
                created_by_user_id=item["user_id"],
                created_at=item["created_at"],
                lead_name=results_data.get("lead", {}).get("name"),
                company_name=results_data.get("lead", {}).get("partner_name")
            ))

        return shared_analyses

    except Exception as exc:
        logger.error(f"Error fetching shared analyses: {exc}")
        raise HTTPException(status_code=500, detail=str(exc))


@app.put("/re-engage/analyses/{analysis_id}")
def update_shared_analysis(
    analysis_id: str,
    payload: SaveAnalysisRequest,
    current_user: Dict[str, Any] = Depends(get_current_user)
) -> SavedAnalysisResponse:
    """Update a shared analysis (only by the creator)."""
    if not supabase:
        raise HTTPException(status_code=503, detail="Database not available")

    try:
        # Update only if user is the creator
        update_data = {
            "results": payload.analysis_data,
            "updated_at": datetime.utcnow().isoformat()
        }
        if payload.title:
            update_data["title"] = payload.title

        result = supabase.client.table("analysis_cache")\
            .update(update_data)\
            .eq("id", analysis_id)\
            .eq("user_id", current_user["id"])\
            .execute()

        if result.data:
            return SavedAnalysisResponse(
                success=True,
                analysis_id=analysis_id,
                message="Analysis updated successfully"
            )
        else:
            raise HTTPException(status_code=404, detail="Analysis not found or access denied")

    except Exception as exc:
        logger.error(f"Error updating analysis {analysis_id}: {exc}")
        raise HTTPException(status_code=500, detail=str(exc))


@app.delete("/re-engage/analyses/{analysis_id}")
def delete_shared_analysis(
    analysis_id: str,
    current_user: Dict[str, Any] = Depends(get_current_user)
) -> Dict[str, Any]:
    """Delete a shared analysis (only by the creator)."""
    if not supabase:
        raise HTTPException(status_code=503, detail="Database not available")

    try:
        # Delete only if user is the creator
        result = supabase.client.table("analysis_cache")\
            .delete()\
            .eq("id", analysis_id)\
            .eq("user_id", current_user["id"])\
            .execute()

        if result.data:
            return {"success": True, "message": "Analysis deleted"}
        else:
            raise HTTPException(status_code=404, detail="Analysis not found or access denied")

    except Exception as exc:
        logger.error(f"Error deleting analysis {analysis_id}: {exc}")
        raise HTTPException(status_code=500, detail=str(exc))


class LostLeadDraftRequest(BaseModel):
    """Request to generate re-engagement email draft for lost lead."""
    lead_data: Dict[str, Any]
    analysis_data: Dict[str, Any]


@app.post("/lost-leads/generate-draft")
def generate_lost_lead_draft(request: LostLeadDraftRequest):
    """Generate a re-engagement email draft for a lost lead."""
    try:
        from modules.llm_client import LLMClient

        llm = LLMClient()

        # Extract lead information
        lead = request.lead_data
        analysis = request.analysis_data.get("analysis", {})

        # Build context for email generation
        lead_name = lead.get("partner_name") or lead.get("contact_name") or lead.get("name", "there")
        company_name = lead.get("partner_name", "")
        lost_reason = lead.get("lost_reason", "Unknown reason")

        # Get analysis insights
        key_insights = analysis.get("key_insights", [])
        recommended_actions = analysis.get("recommended_actions", [])
        summary = analysis.get("summary", "")

        # Create prompt for draft generation
        prompt = f"""You are a sales professional crafting a re-engagement email to a lost lead.

Lead Information:
- Name: {lead_name}
- Company: {company_name}
- Lost Reason: {lost_reason}

Analysis Summary:
{summary}

Key Insights:
{chr(10).join(f"- {insight}" for insight in key_insights[:5]) if key_insights else "No specific insights"}

Recommended Actions:
{chr(10).join(f"- {action}" for action in recommended_actions[:3]) if recommended_actions else "Standard follow-up"}

Please write a professional, personalized re-engagement email that:
1. Acknowledges the previous interaction respectfully
2. Offers genuine value or addresses their concerns
3. Proposes a specific next step or meeting
4. Maintains a warm, non-pushy tone
5. Is concise (3-4 paragraphs maximum)
6. IMPORTANT: Do NOT use em dashes (—). Use regular hyphens (-) or commas instead.
7. IMPORTANT: Do NOT include any signature block (no "Best regards", "Sincerely", "Best", etc.). The user's Outlook signature will be added automatically.

Return your response in JSON format with two fields:
{{
  "subject": "A compelling, personalized subject line (max 60 characters)",
  "body": "The email body text"
}}

The subject line should be attention-grabbing but professional, specific to this lead's situation, and NOT generic."""

        messages = [
            {"role": "system", "content": "You are an expert sales professional writing re-engagement emails to lost leads. Always respond in valid JSON format."},
            {"role": "user", "content": prompt}
        ]

        result = llm.chat_completion_json(messages, max_tokens=6000, temperature=0.7)

        if not result or not result.get("body") or not result.get("subject"):
            raise HTTPException(status_code=500, detail="Failed to generate draft email")

        return {
            "success": True,
            "draft": result["body"],
            "subject_suggestion": result["subject"]
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating lost lead draft: {e}")
        raise HTTPException(status_code=500, detail=str(e))


class SendLostLeadEmailRequest(BaseModel):
    """Request to send re-engagement email for lost lead."""
    lead_id: int
    subject: str
    body: str
    cc: List[str] = []


@app.post("/lost-leads/send-email")
def send_lost_lead_email(
    request: SendLostLeadEmailRequest,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """Send a re-engagement email to a lost lead."""
    try:
        from modules.outlook_client import OutlookClient
        from modules.odoo_client import OdooClient

        # Get lead details from Odoo
        config = Config()
        odoo = OdooClient(config)
        if not odoo.connect():
            raise HTTPException(status_code=500, detail="Failed to connect to Odoo")

        # Fetch the lead to get email
        lead = odoo._call_kw(
            'crm.lead', 'search_read',
            [[['id', '=', request.lead_id]]],
            {'fields': ['email_from', 'partner_name', 'contact_name'], 'limit': 1}
        )

        if not lead:
            raise HTTPException(status_code=404, detail="Lead not found")

        lead_email = lead[0].get('email_from')
        if not lead_email:
            raise HTTPException(status_code=400, detail="Lead has no email address")

        # Get Outlook tokens
        outlook = OutlookClient()
        user_identifier = str(current_user["id"])
        tokens = outlook.get_user_auth_tokens(user_identifier)

        if not tokens or 'access_token' not in tokens:
            raise HTTPException(
                status_code=401,
                detail="No authenticated Outlook account found. Please authenticate in Settings."
            )

        # Send the email
        success = outlook.send_email_with_attachment(
            access_token=tokens['access_token'],
            to=[lead_email],
            subject=request.subject,
            body=request.body.replace('\n', '<br>'),
            cc=request.cc,
            attachment_path=None,  # No attachment for re-engagement emails
            attachment_name=None
        )

        if not success:
            raise HTTPException(status_code=500, detail="Failed to send email")

        logger.info(f"Re-engagement email sent to {lead_email} for lead {request.lead_id}")

        return {
            "success": True,
            "message": f"Email sent successfully to {lead_email}"
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error sending re-engagement email: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/lost-leads/generate-report")
def generate_lost_leads_report(
    limit: int = 50,
    salesperson: Optional[str] = None,
    type_filter: Optional[str] = None,
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
    include_pattern_analysis: bool = False,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """
    Generate comprehensive lost leads report with statistics and analysis.

    Args:
        limit: Max number of lost leads to analyze
        salesperson: Filter by salesperson name
        type_filter: Filter by 'lead' or 'opportunity'
        date_from: Start date in YYYY-MM-DD format (when marked as lost)
        date_to: End date in YYYY-MM-DD format (when marked as lost)
        include_pattern_analysis: Generate LLM-powered pattern analysis (6 dimensions)

    Returns:
        - Summary statistics (total missed value, average deal size, counts)
        - Reasons analysis (by frequency and by value)
        - Monthly trends (month-by-month breakdown)
        - Top 10 opportunities to re-contact (scored by value, recency, reason)
        - Optional: Pattern analysis (6 LLM-powered insights)
    """
    analyzer = get_lost_lead_analyzer()

    try:
        logger.info(f"Generating lost leads report: limit={limit}, salesperson={salesperson}, type={type_filter}, date_from={date_from}, date_to={date_to}, pattern_analysis={include_pattern_analysis}")
        report = analyzer.generate_lost_leads_report(
            limit=limit,
            salesperson_name=salesperson,
            type_filter=type_filter,
            date_from=date_from,
            date_to=date_to
        )

        # Optionally add LLM pattern analysis
        if include_pattern_analysis and report.get("all_lost_leads"):
            logger.info("Generating LLM pattern analysis...")
            pattern_analysis = analyzer.generate_pattern_analysis(
                lost_leads=report["all_lost_leads"],
                limit=100
            )
            report["pattern_analysis"] = pattern_analysis

        return {
            "success": True,
            "report": report
        }

    except Exception as e:
        logger.error(f"Error generating lost leads report: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# PROPOSAL FOLLOW-UP ENDPOINTS
# ============================================================================

class ProposalFollowupSettings(BaseModel):
    """Settings for proposal follow-up analysis."""
    days_back: int = 7
    no_response_days: int = 3


class ProposalFollowupSummary(BaseModel):
    """Summary counts for proposal follow-ups."""
    unanswered_count: int
    pending_proposals_count: int
    filtered_count: Optional[int] = 0
    total_count: int
    days_back: int
    no_response_days: int
    last_updated: Optional[str] = None


class ProposalFollowupThread(BaseModel):
    """Single email thread needing follow-up."""
    conversation_id: str
    external_email: str
    subject: str
    days_waiting: int
    last_contact_date: Optional[str] = None
    proposal_date: Optional[str] = None
    odoo_lead: Optional[Dict[str, Any]] = None
    analysis: Optional[Dict[str, Any]] = None
    classification: Optional[Dict[str, Any]] = None  # AI classification: is_lead, confidence, category
    is_favorited: Optional[bool] = None


class ProposalFollowupResponse(BaseModel):
    """Response containing all proposal follow-up data."""
    summary: ProposalFollowupSummary
    unanswered: List[ProposalFollowupThread]
    pending_proposals: List[ProposalFollowupThread]
    filtered: Optional[List[ProposalFollowupThread]] = []


@app.get("/proposal-followups", response_model=ProposalFollowupResponse)
def get_proposal_followups(
    days_back: int = 3,
    no_response_days: int = 3,
    engage_email: str = "automated.response@prezlab.com",
    force_refresh: bool = False,
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: SupabaseDatabase = Depends(get_supabase_database)
):
    """
    Get proposal follow-up analysis from engage inbox.

    Args:
        days_back: Number of days to look back for emails (default: 3)
        no_response_days: Days threshold for "no response" (default: 3)
        engage_email: Email of the engage monitoring account
        force_refresh: Force refresh analysis instead of using cache
        current_user: Current authenticated user
        db: Database instance

    Returns:
        Summary and categorized threads needing follow-up with last_updated timestamp
    """
    global proposal_followups_cache
    user_id = current_user.get("id")

    # Determine cache duration based on days_back
    # 90-day analysis: cache for 90 days
    # 7-day (weekly): cache for 7 days
    # Others: cache for 1 day
    if days_back >= 90:
        cache_duration_days = 90
    elif days_back >= 7:
        cache_duration_days = 7
    else:
        cache_duration_days = 1

    # Parameters for caching
    cache_params = {
        "days_back": days_back,
        "no_response_days": no_response_days,
        "engage_email": engage_email
    }

    # Get completed thread IDs with timestamps to check for new activity
    completed_threads_map = {}
    try:
        completed_threads_map = db.get_completed_followups_with_timestamps()
    except Exception as e:
        logger.warning(f"Failed to get completed followups: {e}")

    # Will be populated with threads that should stay completed
    completed_thread_ids = set()

    # Get favorited thread IDs
    favorited_thread_ids = set()
    try:
        favorited_thread_ids = set(db.get_favorited_followups())
    except Exception as e:
        logger.warning(f"Failed to get favorited followups: {e}")

    # Try to get from Supabase cache first (if available)
    if not force_refresh and supabase.is_connected():
        try:
            cached_data = supabase.get_cached_analysis(
                user_id=user_id,
                analysis_type="proposal_followups",
                parameters=cache_params
            )

            if cached_data:
                logger.info(f"✅ Returning cached proposal follow-ups from Supabase for user {user_id}")
                # Filter out completed threads and add favorited flag
                unanswered = []
                for thread in cached_data.get("unanswered", []):
                    if thread.get("conversation_id") not in completed_thread_ids:
                        thread["is_favorited"] = thread.get("conversation_id") in favorited_thread_ids
                        unanswered.append(thread)

                pending_proposals = []
                for thread in cached_data.get("pending_proposals", []):
                    if thread.get("conversation_id") not in completed_thread_ids:
                        thread["is_favorited"] = thread.get("conversation_id") in favorited_thread_ids
                        pending_proposals.append(thread)

                # Sort: favorited threads first, then by days_waiting descending
                unanswered.sort(key=lambda x: (not x.get("is_favorited", False), -x.get("days_waiting", 0)))
                pending_proposals.sort(key=lambda x: (not x.get("is_favorited", False), -x.get("days_waiting", 0)))

                # Update summary counts
                summary = cached_data["summary"].copy()
                summary["unanswered_count"] = len(unanswered)
                summary["pending_proposals_count"] = len(pending_proposals)
                summary["total_count"] = len(unanswered) + len(pending_proposals)

                # Convert cached data back to response model
                response = ProposalFollowupResponse(
                    summary=ProposalFollowupSummary(**summary),
                    unanswered=[ProposalFollowupThread(**thread) for thread in unanswered],
                    pending_proposals=[ProposalFollowupThread(**thread) for thread in pending_proposals],
                    filtered=[ProposalFollowupThread(**thread) for thread in cached_data.get("filtered", [])]
                )
                return response
        except Exception as e:
            logger.warning(f"Failed to get from Supabase cache: {e}. Falling back to in-memory cache.")

    # Fallback to in-memory cache if Supabase unavailable
    if not force_refresh and proposal_followups_cache["data"] and proposal_followups_cache["params"] == cache_params:
        logger.info("Returning cached proposal follow-ups analysis from memory")
        cached_response = proposal_followups_cache["data"]
        # Filter out completed threads and add favorited flag
        cached_response.unanswered = [thread for thread in cached_response.unanswered
                                     if thread.conversation_id not in completed_thread_ids]
        for thread in cached_response.unanswered:
            thread.is_favorited = thread.conversation_id in favorited_thread_ids

        cached_response.pending_proposals = [thread for thread in cached_response.pending_proposals
                                            if thread.conversation_id not in completed_thread_ids]
        for thread in cached_response.pending_proposals:
            thread.is_favorited = thread.conversation_id in favorited_thread_ids

        # Sort: favorited first, then by days_waiting
        cached_response.unanswered.sort(key=lambda x: (not x.is_favorited, -x.days_waiting))
        cached_response.pending_proposals.sort(key=lambda x: (not x.is_favorited, -x.days_waiting))

        # Update counts
        cached_response.summary.unanswered_count = len(cached_response.unanswered)
        cached_response.summary.pending_proposals_count = len(cached_response.pending_proposals)
        cached_response.summary.total_count = len(cached_response.unanswered) + len(cached_response.pending_proposals)
        cached_response.summary.last_updated = proposal_followups_cache["timestamp"]
        return cached_response

    try:
        logger.info(f"Running new proposal follow-ups analysis (days_back={days_back}, no_response_days={no_response_days})")
        analyzer = ProposalFollowupAnalyzer()
        # Use SYSTEM_ prefix for system email tokens
        system_identifier = f"SYSTEM_{engage_email}"
        result = analyzer.get_proposal_followups(
            user_identifier=system_identifier,
            days_back=days_back,
            no_response_days=no_response_days
        )

        # Check for new activity in completed threads and reopen if needed
        from datetime import datetime, timezone
        for thread in result["unanswered"] + result["pending_proposals"]:
            conv_id = thread.get("conversation_id")

            # If thread was completed, check if there's new activity
            if conv_id in completed_threads_map:
                completed_at_str = completed_threads_map[conv_id]
                thread_last_email_str = thread.get("last_email_date")

                if thread_last_email_str and completed_at_str:
                    try:
                        # Parse timestamps
                        completed_at = datetime.fromisoformat(completed_at_str.replace('Z', '+00:00'))
                        last_email = datetime.fromisoformat(thread_last_email_str.replace('Z', '+00:00'))

                        # If last email is AFTER completion, reopen the thread
                        if last_email > completed_at:
                            logger.info(f"Reopening thread {conv_id}: new email at {last_email} after completion at {completed_at}")
                            db.reopen_completed_followup(conv_id)
                            # Don't add to completed_thread_ids, so it appears in results
                        else:
                            # No new activity, keep it completed
                            completed_thread_ids.add(conv_id)
                    except Exception as parse_error:
                        logger.warning(f"Error parsing dates for thread {conv_id}: {parse_error}")
                        # If we can't parse, keep it completed to be safe
                        completed_thread_ids.add(conv_id)
                else:
                    # Missing date info, keep it completed
                    completed_thread_ids.add(conv_id)

        # Filter out completed threads from fresh results and add favorited flag
        unanswered = []
        for thread in result["unanswered"]:
            if thread.get("conversation_id") not in completed_thread_ids:
                thread["is_favorited"] = thread.get("conversation_id") in favorited_thread_ids
                unanswered.append(thread)

        pending_proposals = []
        for thread in result["pending_proposals"]:
            if thread.get("conversation_id") not in completed_thread_ids:
                thread["is_favorited"] = thread.get("conversation_id") in favorited_thread_ids
                pending_proposals.append(thread)

        # Sort: favorited first, then by days_waiting
        unanswered.sort(key=lambda x: (not x.get("is_favorited", False), -x.get("days_waiting", 0)))
        pending_proposals.sort(key=lambda x: (not x.get("is_favorited", False), -x.get("days_waiting", 0)))

        # Update summary counts
        result["summary"]["unanswered_count"] = len(unanswered)
        result["summary"]["pending_proposals_count"] = len(pending_proposals)
        result["summary"]["total_count"] = len(unanswered) + len(pending_proposals)

        # Create timestamp
        timestamp = datetime.utcnow().isoformat() + "Z"

        response = ProposalFollowupResponse(
            summary=ProposalFollowupSummary(**result["summary"], last_updated=timestamp),
            unanswered=[ProposalFollowupThread(**thread) for thread in unanswered],
            pending_proposals=[ProposalFollowupThread(**thread) for thread in pending_proposals],
            filtered=[ProposalFollowupThread(**thread) for thread in result.get("filtered", [])]
        )

        # Convert response to dict for caching
        response_dict = {
            "summary": response.summary.dict(),
            "unanswered": [thread.dict() for thread in response.unanswered],
            "pending_proposals": [thread.dict() for thread in response.pending_proposals],
            "filtered": [thread.dict() for thread in response.filtered]
        }

        # Save to Supabase cache
        if supabase.is_connected():
            try:
                supabase.save_analysis_cache(
                    user_id=user_id,
                    analysis_type="proposal_followups",
                    parameters=cache_params,
                    results=response_dict,
                    cache_duration_days=cache_duration_days
                )
                logger.info(f"✅ Saved analysis to Supabase cache (expires in {cache_duration_days} days)")
            except Exception as e:
                logger.warning(f"Failed to save to Supabase cache: {e}")

        # Also save to in-memory cache as fallback
        proposal_followups_cache = {
            "data": response,
            "timestamp": timestamp,
            "params": cache_params
        }

        return response
    except ValueError as e:
        raise HTTPException(status_code=401, detail=str(e))
    except Exception as e:
        logger.error(f"Error getting proposal follow-ups: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/proposal-followups/analyze-thread")
def analyze_followup_thread(thread_data: Dict[str, Any]):
    """
    Analyze a specific email thread and generate follow-up draft.

    Args:
        thread_data: Thread data with conversation_id and thread messages

    Returns:
        Analysis with summary, sentiment, urgency, and draft email
    """
    try:
        analyzer = ProposalFollowupAnalyzer()
        analysis = analyzer.analyze_thread_with_llm(thread_data)
        return analysis
    except Exception as e:
        logger.error(f"Error analyzing thread: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# FOLLOW-UP REPORTS AND MANAGEMENT ENDPOINTS
# ============================================================================

class SavedReport(BaseModel):
    id: str  # UUID from Supabase
    report_type: str
    report_period: str
    created_at: str
    result: Dict[str, Any]
    parameters: Dict[str, Any]


class SavedReportsResponse(BaseModel):
    count: int
    reports: List[SavedReport]


class GenerateReportRequest(BaseModel):
    report_type: Literal["90day", "monthly", "weekly", "complete"]
    days_back: Optional[int] = None
    no_response_days: int = 3
    engage_email: str = "automated.response@prezlab.com"


class MarkCompleteRequest(BaseModel):
    thread_id: str
    conversation_id: str
    notes: Optional[str] = None


class GenerateDraftRequest(BaseModel):
    thread_data: Dict[str, Any]


class RefineDraftRequest(BaseModel):
    current_draft: str
    edit_prompt: str
    thread_data: Dict[str, Any]


class SendFollowupEmailRequest(BaseModel):
    conversation_id: str
    draft_body: str
    subject: str
    reply_to_message_id: Optional[str] = None


@app.get("/proposal-followups/reports", response_model=SavedReportsResponse)
def get_saved_reports(
    report_type: Optional[str] = None,
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: SupabaseDatabase = Depends(get_supabase_database)
):
    """Get all saved follow-up reports."""
    try:
        reports = db.get_saved_reports(
            analysis_type="proposal_followups",
            report_type=report_type
        )

        return SavedReportsResponse(
            count=len(reports),
            reports=[SavedReport(**report) for report in reports]
        )
    except Exception as e:
        logger.error(f"Error fetching saved reports: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/proposal-followups/reports/{report_id}")
def delete_saved_report(
    report_id: str,  # UUID string from Supabase
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: SupabaseDatabase = Depends(get_supabase_database)
):
    """Delete a saved follow-up report."""
    try:
        success = db.delete_report(report_id)
        if not success:
            raise HTTPException(status_code=500, detail="Failed to delete report")

        return {"success": True, "message": "Report deleted successfully"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting report {report_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/proposal-followups/reports/{report_id}/export")
def export_report_pdf(
    report_id: str,
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: SupabaseDatabase = Depends(get_supabase_database)
):
    """Export a follow-up report as PDF with executive summary."""
    try:
        # Get report from database
        report = db.get_report(report_id)
        if not report:
            raise HTTPException(status_code=404, detail="Report not found")

        # Parse report data
        report_data = report.get("report_data", {})

        # Combine unanswered and pending_proposals into followups list
        unanswered = report_data.get("unanswered", [])
        pending_proposals = report_data.get("pending_proposals", [])
        followups = unanswered + pending_proposals

        # Generate executive summary using OpenAI
        executive_summary = generate_executive_summary(followups, report.get("report_type", ""))

        # Create PDF
        pdf_buffer = BytesIO()
        doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1f2937'),
            spaceAfter=30,
            alignment=TA_CENTER
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#374151'),
            spaceAfter=12,
            spaceBefore=20
        )

        # Title
        story.append(Paragraph("Proposal Follow-up Report", title_style))
        story.append(Spacer(1, 0.3*inch))

        # Report metadata
        report_type = report.get("report_type", "").replace("_", " ").title()
        report_period = report.get("report_period", "")
        created_at = report.get("created_at", "")

        metadata_data = [
            ["Report Type:", report_type],
            ["Period:", report_period],
            ["Generated:", created_at[:10] if created_at else ""],
            ["Total Follow-ups:", str(len(followups))]
        ]

        metadata_table = Table(metadata_data, colWidths=[2*inch, 4*inch])
        metadata_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#374151')),
            ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
            ('ALIGN', (1, 0), (1, -1), 'LEFT'),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))

        story.append(metadata_table)
        story.append(Spacer(1, 0.5*inch))

        # Executive Summary
        story.append(Paragraph("Executive Summary", heading_style))
        story.append(Paragraph(executive_summary, styles['BodyText']))
        story.append(Spacer(1, 0.3*inch))

        # Statistics
        story.append(Paragraph("Key Statistics", heading_style))

        # Calculate stats
        urgent_count = len([f for f in followups if f.get("days_since_proposal", 0) > 14])
        recent_count = len([f for f in followups if f.get("days_since_proposal", 0) <= 7])

        stats_data = [
            ["Total Follow-ups Needed:", str(len(followups))],
            ["Urgent (>14 days):", str(urgent_count)],
            ["Recent (≤7 days):", str(recent_count)],
        ]

        stats_table = Table(stats_data, colWidths=[3*inch, 2*inch])
        stats_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#374151')),
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f3f4f6')),
            ('ALIGN', (0, 0), (0, -1), 'LEFT'),
            ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('LEFTPADDING', (0, 0), (-1, -1), 12),
            ('RIGHTPADDING', (0, 0), (-1, -1), 12),
        ]))

        story.append(stats_table)
        story.append(Spacer(1, 0.5*inch))

        # Follow-up threads
        story.append(PageBreak())
        story.append(Paragraph("Follow-up Threads", heading_style))
        story.append(Spacer(1, 0.2*inch))

        for idx, followup in enumerate(followups[:50], 1):  # Limit to first 50 for PDF size
            # Thread header
            subject = followup.get("subject", "No subject")
            external_email = followup.get("external_email", "Unknown")
            days_since = followup.get("days_since_proposal", 0)

            thread_title = f"{idx}. {subject}"
            story.append(Paragraph(thread_title, styles['Heading3']))

            # Thread details
            details_data = [
                ["Contact:", external_email],
                ["Days Since Proposal:", str(days_since)],
            ]

            details_table = Table(details_data, colWidths=[1.5*inch, 4.5*inch])
            details_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#4b5563')),
                ('TOPPADDING', (0, 0), (-1, -1), 3),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
            ]))

            story.append(details_table)
            story.append(Spacer(1, 0.2*inch))

        if len(followups) > 50:
            story.append(Paragraph(
                f"<i>Note: Only showing first 50 of {len(followups)} total follow-ups. View full report in the app.</i>",
                styles['BodyText']
            ))

        # Build PDF
        doc.build(story)
        pdf_buffer.seek(0)

        # Save to temp file and return
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            tmp_file.write(pdf_buffer.getvalue())
            tmp_path = tmp_file.name

        return FileResponse(
            tmp_path,
            media_type='application/pdf',
            filename=f'follow-up-report-{report_id}.pdf'
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error exporting report {report_id}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


def generate_executive_summary(followups: List[Dict], report_type: str) -> str:
    """Generate an executive summary of the follow-ups using OpenAI."""
    try:
        from openai import OpenAI
        client = OpenAI()

        # Prepare summary of followups for AI
        urgent = [f for f in followups if f.get("days_since_proposal", 0) > 14]
        recent = [f for f in followups if f.get("days_since_proposal", 0) <= 7]

        prompt = f"""Based on this proposal follow-up report, write a concise executive summary (3-4 paragraphs):

Report Type: {report_type}
Total Follow-ups Needed: {len(followups)}
Urgent (>14 days): {len(urgent)}
Recent (≤7 days): {len(recent)}

Key patterns to analyze:
- Average days since proposal: {sum(f.get('days_since_proposal', 0) for f in followups) / len(followups) if followups else 0:.1f}
- Distribution of follow-up urgency

Provide:
1. Overview of the follow-up situation
2. Key concerns and priorities
3. Recommended actions

Keep it professional and actionable."""

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a sales operations analyst providing executive summaries."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=500
        )

        return response.choices[0].message.content.strip()

    except Exception as e:
        logger.error(f"Error generating executive summary: {e}")
        # Return fallback summary
        return f"""This report contains {len(followups)} proposal follow-ups requiring attention.

Of these, {len([f for f in followups if f.get('days_since_proposal', 0) > 14])} are considered urgent (over 14 days since proposal),
while {len([f for f in followups if f.get('days_since_proposal', 0) <= 7])} are recent (within 7 days).

Priority should be given to the urgent follow-ups to maintain engagement and move deals forward.
The team should systematically work through these opportunities to maximize conversion rates."""


@app.get("/outlook/conversation/{conversation_id}")
def get_conversation_thread(
    conversation_id: str,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """Fetch all messages in an Outlook conversation thread."""
    try:
        outlook = get_outlook_client()
        user_identifier = str(current_user["id"])

        # Get user's Outlook tokens
        tokens = outlook.get_user_auth_tokens(user_identifier)
        if not tokens or 'access_token' not in tokens:
            raise HTTPException(
                status_code=401,
                detail="No authenticated Outlook account found. Please authenticate in Settings."
            )

        access_token = tokens['access_token']

        # Try to fetch from shared mailbox first, fall back to personal mailbox if access denied
        messages = []
        try:
            # First try shared mailbox (engage email)
            logger.info(f"Attempting to fetch conversation from shared mailbox: automated.response@prezlab.com")
            messages = outlook.get_conversation_messages(
                access_token,
                conversation_id,
                shared_mailbox="automated.response@prezlab.com"
            )
        except Exception as e:
            if "403" in str(e) or "Forbidden" in str(e) or "Access" in str(e):
                # User doesn't have permission to shared mailbox, try personal mailbox
                logger.info(f"Shared mailbox access denied, falling back to personal mailbox")
                messages = outlook.get_conversation_messages(
                    access_token,
                    conversation_id,
                    shared_mailbox=None  # Use personal mailbox
                )
            else:
                # Re-raise other errors
                raise

        if not messages:
            return {
                "conversation_id": conversation_id,
                "messages": [],
                "count": 0
            }

        # Format messages for frontend
        formatted_messages = []
        for msg in messages:
            from_data = msg.get("from", {}).get("emailAddress", {})
            to_recipients = msg.get("toRecipients", [])
            cc_recipients = msg.get("ccRecipients", [])

            formatted_messages.append({
                "id": msg.get("id"),
                "subject": msg.get("subject", ""),
                "from": {
                    "name": from_data.get("name", "Unknown"),
                    "email": from_data.get("address", "")
                },
                "to": [{"name": r.get("emailAddress", {}).get("name", ""), "email": r.get("emailAddress", {}).get("address", "")} for r in to_recipients],
                "cc": [{"name": r.get("emailAddress", {}).get("name", ""), "email": r.get("emailAddress", {}).get("address", "")} for r in cc_recipients],
                "receivedDateTime": msg.get("receivedDateTime", ""),
                "body": msg.get("body", {}).get("content", ""),
                "bodyPreview": msg.get("bodyPreview", ""),
                "hasAttachments": msg.get("hasAttachments", False),
                "webLink": msg.get("webLink", "")
            })

        return {
            "conversation_id": conversation_id,
            "messages": formatted_messages,
            "count": len(formatted_messages)
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching conversation thread: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/proposal-followups/reports/generate")
def generate_saved_report(
    request: GenerateReportRequest,
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: SupabaseDatabase = Depends(get_supabase_database)
):
    """Generate a new follow-up report and save it."""
    try:
        user_id = current_user.get("id")
        user_identifier = current_user.get("email")

        # Determine days_back based on report type
        if request.report_type == "complete":
            days_back = 365  # Analyze all emails from the past year
            report_period = datetime.now().strftime("%Y")
        elif request.report_type == "90day":
            days_back = 90
            report_period = datetime.now().strftime("%Y-Q%q")
        elif request.report_type == "monthly":
            days_back = 30
            report_period = datetime.now().strftime("%Y-%m")
        else:  # weekly
            days_back = 7
            report_period = datetime.now().strftime("%Y-W%W")

        # Generate the analysis
        logger.info(f"Starting report generation for {request.report_type} (days_back={days_back})")
        analyzer = ProposalFollowupAnalyzer()
        # Use SYSTEM_ prefix for system email tokens
        system_identifier = f"SYSTEM_{request.engage_email}"
        result = analyzer.get_proposal_followups(
            user_identifier=system_identifier,
            days_back=days_back,
            no_response_days=request.no_response_days
        )
        logger.info(f"Analysis completed. Found {result.get('summary', {}).get('total_count', 0)} follow-ups")

        # Save as shared report
        logger.info(f"Attempting to save report to database for user {user_id}")
        report_id = db.save_report(
            user_id=user_id,
            analysis_type="proposal_followups",
            report_type=request.report_type,
            report_period=report_period,
            result=result,
            parameters={
                "days_back": days_back,
                "no_response_days": request.no_response_days,
                "engage_email": request.engage_email
            },
            is_shared=True
        )
        logger.info(f"Report saved successfully with ID: {report_id}")

        return {
            "success": True,
            "report_id": report_id,
            "report_type": request.report_type,
            "report_period": report_period
        }

    except Exception as e:
        logger.error(f"Error generating report: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/proposal-followups/daily-digest/send")
def send_daily_digest(
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: SupabaseDatabase = Depends(get_supabase_database)
):
    """Send daily digest to Teams based on most recent complete report."""
    try:
        from modules.daily_digest_formatter import DailyDigestFormatter
        from modules.teams_messenger import TeamsMessenger
        from modules.email_token_store import EmailTokenStore

        # Get Microsoft access token
        token_store = EmailTokenStore()
        user_id = str(current_user.get("id"))
        user_email = current_user.get("email")

        logger.info(f"Getting tokens for user ID: {user_id} (email: {user_email})")
        tokens = token_store.get_tokens(user_id)

        if not tokens or not isinstance(tokens, dict):
            logger.error(f"No valid tokens found for {user_email}")
            raise HTTPException(
                status_code=401,
                detail="No Microsoft authentication found. Please connect your Microsoft account in Settings."
            )

        access_token = tokens.get("access_token")
        if not access_token:
            logger.error(f"Token data found but no access_token field for {user_email}")
            raise HTTPException(
                status_code=401,
                detail="Microsoft token is invalid. Please reconnect your Microsoft account in Settings."
            )

        # Get the most recent complete report
        reports = db.get_saved_reports(
            analysis_type="proposal_followups",
            report_type="complete"
        )

        if not reports:
            raise HTTPException(
                status_code=404,
                detail="No complete reports found. Please generate a complete report first."
            )

        latest_report = reports[0]  # Reports are sorted by created_at desc
        report_data = latest_report.get('result', {})

        # Format the daily digest
        digest_html = DailyDigestFormatter.format_digest(report_data)

        # Send to Teams
        teams_chat_id = '19:1d7fae90086342a49e12a433576697c7@thread.v2'
        teams = TeamsMessenger(access_token)
        teams.send_message_to_chat(teams_chat_id, digest_html)

        logger.info(f"Daily digest sent successfully to Teams chat {teams_chat_id}")

        return {
            "success": True,
            "message": "Daily digest sent to Teams",
            "report_date": latest_report.get('created_at'),
            "total_followups": report_data.get('summary', {}).get('total_count', 0)
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error sending daily digest: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/proposal-followups/{thread_id}/mark-complete")
def mark_followup_complete(
    thread_id: str,
    request: MarkCompleteRequest,
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: SupabaseDatabase = Depends(get_supabase_database)
):
    """Mark a follow-up as completed and remove from saved reports."""
    try:
        user_id = current_user.get("id")

        result = db.mark_followup_complete(
            thread_id=request.thread_id,
            conversation_id=request.conversation_id,
            user_id=user_id,
            completion_method="manual_marked",
            notes=request.notes
        )

        # Update all saved reports to remove this thread
        try:
            # Get all saved proposal followup reports
            reports = supabase.client.table("analysis_cache")\
                .select("*")\
                .eq("analysis_type", "proposal_followups")\
                .eq("is_shared", True)\
                .execute()

            if reports.data:
                for report in reports.data:
                    results = report.get("results", {})

                    # Skip if results is not a dict (corrupted data)
                    if not isinstance(results, dict):
                        logger.warning(f"Skipping report {report.get('id')} - results is not a dict")
                        continue

                    updated = False

                    # Remove from unanswered emails
                    if "unanswered" in results and isinstance(results["unanswered"], list):
                        original_count = len(results["unanswered"])
                        results["unanswered"] = [
                            t for t in results["unanswered"]
                            if t.get("conversation_id") != request.conversation_id
                        ]
                        if len(results["unanswered"]) < original_count:
                            updated = True

                    # Remove from pending proposals
                    if "pending_proposals" in results and isinstance(results["pending_proposals"], list):
                        original_count = len(results["pending_proposals"])
                        results["pending_proposals"] = [
                            t for t in results["pending_proposals"]
                            if t.get("conversation_id") != request.conversation_id
                        ]
                        if len(results["pending_proposals"]) < original_count:
                            updated = True

                    # Update summary counts if summary exists and is a dict
                    if "summary" in results and isinstance(results["summary"], dict):
                        results["summary"]["unanswered_count"] = len(results.get("unanswered", []))
                        results["summary"]["pending_count"] = len(results.get("pending_proposals", []))
                        results["summary"]["total_count"] = results["summary"]["unanswered_count"] + results["summary"]["pending_count"]

                    # Save updated report if changed
                    if updated:
                        supabase.client.table("analysis_cache")\
                            .update({"results": results})\
                            .eq("id", report["id"])\
                            .execute()
                        logger.info(f"Updated report {report['id']} to remove completed thread {request.conversation_id}")

        except Exception as update_error:
            logger.error(f"Error updating saved reports: {update_error}")
            # Don't fail the completion if report update fails

        return {"success": True, "completion": result}

    except Exception as e:
        logger.error(f"Error marking follow-up complete: {e}")

        # Check if it's a duplicate key error (already marked complete)
        error_str = str(e)
        if "duplicate key" in error_str or "23505" in error_str or "already exists" in error_str:
            # Thread already marked complete - still update reports to remove it
            logger.info(f"Thread already marked complete, updating reports to remove it")

            # Update all saved reports to remove this thread
            try:
                reports = supabase.client.table("analysis_cache")\
                    .select("*")\
                    .eq("analysis_type", "proposal_followups")\
                    .eq("is_shared", True)\
                    .execute()

                if reports.data:
                    for report in reports.data:
                        results = report.get("results", {})

                        # Skip if results is not a dict (corrupted data)
                        if not isinstance(results, dict):
                            logger.warning(f"Skipping report {report.get('id')} - results is not a dict")
                            continue

                        updated = False

                        # Remove from unanswered emails
                        if "unanswered" in results and isinstance(results["unanswered"], list):
                            original_count = len(results["unanswered"])
                            results["unanswered"] = [
                                t for t in results["unanswered"]
                                if t.get("conversation_id") != request.conversation_id
                            ]
                            if len(results["unanswered"]) < original_count:
                                updated = True

                        # Remove from pending proposals
                        if "pending_proposals" in results and isinstance(results["pending_proposals"], list):
                            original_count = len(results["pending_proposals"])
                            results["pending_proposals"] = [
                                t for t in results["pending_proposals"]
                                if t.get("conversation_id") != request.conversation_id
                            ]
                            if len(results["pending_proposals"]) < original_count:
                                updated = True

                        # Update summary counts if summary exists and is a dict
                        if "summary" in results and isinstance(results["summary"], dict):
                            results["summary"]["unanswered_count"] = len(results.get("unanswered", []))
                            results["summary"]["pending_count"] = len(results.get("pending_proposals", []))
                            results["summary"]["total_count"] = results["summary"]["unanswered_count"] + results["summary"]["pending_count"]

                        # Save updated report if changed
                        if updated:
                            supabase.client.table("analysis_cache")\
                                .update({"results": results})\
                                .eq("id", report["id"])\
                                .execute()
                            logger.info(f"Updated report {report['id']} to remove already-completed thread {request.conversation_id}")

            except Exception as update_error:
                logger.error(f"Error updating saved reports for duplicate completion: {update_error}")

            return {
                "success": True,
                "completion": {"already_completed": True},
                "message": "This follow-up was already marked as complete"
            }

        raise HTTPException(status_code=500, detail=str(e))


@app.get("/proposal-followups/completions")
def get_completed_followups(
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: SupabaseDatabase = Depends(get_supabase_database)
):
    """Get list of all completed followup thread IDs."""
    try:
        completions = db.get_completed_followups_with_timestamps()
        # Return as list of objects with conversation_id and completed_at
        return [
            {"conversation_id": conv_id, "completed_at": timestamp}
            for conv_id, timestamp in completions.items()
        ]
    except Exception as e:
        logger.error(f"Error fetching completions: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/proposal-followups/{thread_id}/favorite")
def favorite_followup(
    thread_id: str,
    request: MarkCompleteRequest,  # Reuse same request model (has thread_id and conversation_id)
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: SupabaseDatabase = Depends(get_supabase_database)
):
    """Mark a follow-up thread as favorited."""
    try:
        success = db.favorite_followup(
            thread_id=request.thread_id,
            conversation_id=request.conversation_id
        )

        if not success:
            raise HTTPException(status_code=500, detail="Failed to favorite thread")

        return {"success": True, "message": "Thread favorited successfully"}

    except Exception as e:
        logger.error(f"Error favoriting follow-up: {e}")

        # Check if it's a duplicate key error (already favorited)
        error_str = str(e)
        if "duplicate key" in error_str or "23505" in error_str or "already exists" in error_str:
            return {
                "success": True,
                "message": "This thread is already favorited"
            }

        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/proposal-followups/{thread_id}/favorite")
def unfavorite_followup(
    thread_id: str,
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: SupabaseDatabase = Depends(get_supabase_database)
):
    """Remove favorite from a follow-up thread."""
    try:
        success = db.unfavorite_followup(thread_id=thread_id)

        if not success:
            raise HTTPException(status_code=500, detail="Failed to unfavorite thread")

        return {"success": True, "message": "Thread unfavorited successfully"}

    except Exception as e:
        logger.error(f"Error unfavoriting follow-up: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/proposal-followups/send-to-teams")
def send_report_to_teams(
    request: Dict[str, Any],
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: SupabaseDatabase = Depends(get_supabase_database)
):
    """
    Send a follow-up report summary to a Teams group chat.

    Request body:
    {
        "chat_id": "19:1d7fae90086342a49e12a433576697c7@thread.v2",
        "report_data": {...}  // The report data from generate endpoint
    }
    """
    try:
        chat_id = request.get('chat_id')
        report_data = request.get('report_data')

        if not chat_id or not report_data:
            raise HTTPException(status_code=400, detail="chat_id and report_data are required")

        # Get user's Microsoft access token (for Teams messaging)
        # Teams messages will be sent from the authenticated user
        token_store = EmailTokenStore()
        user_id = str(current_user.get("id"))  # Use user ID, not email
        user_email = current_user.get("email")

        logger.info(f"Attempting to get tokens for user ID: {user_id} (email: {user_email})")
        tokens = token_store.get_tokens(user_id)

        logger.info(f"Tokens retrieved: {tokens is not None}")
        if tokens:
            logger.info(f"Token keys: {list(tokens.keys()) if isinstance(tokens, dict) else 'not a dict'}")

        if not tokens or not isinstance(tokens, dict):
            logger.error(f"No valid tokens found for {user_email}")
            raise HTTPException(
                status_code=401,
                detail="No Microsoft authentication found. Please connect your Microsoft account in Settings."
            )

        access_token = tokens.get("access_token")

        if not access_token:
            logger.error(f"Token data found but no access_token field for {user_email}")
            raise HTTPException(
                status_code=401,
                detail="Microsoft token is invalid. Please reconnect your Microsoft account in Settings."
            )

        # Initialize Teams messenger
        teams = TeamsMessenger(access_token)

        # Format the report as HTML message
        message_html = TeamsMessenger.format_followup_report_summary(report_data)

        # Send to Teams
        result = teams.send_message_to_chat(chat_id, message_html)

        logger.info(f"Successfully sent report to Teams chat {chat_id} for user {current_user.get('email')}")

        return {
            "success": True,
            "message": "Report sent to Teams successfully",
            "teams_message_id": result.get("id")
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error sending report to Teams: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# Weekly Pipeline Performance Endpoints
# ============================================================================

@app.get("/pipeline/weekly-report")
def get_weekly_pipeline_report(
    week_start: Optional[str] = None,
    week_end: Optional[str] = None,
    salesperson_filter: Optional[str] = None,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """
    Generate weekly pipeline performance report.

    Query params:
    - week_start: Start date (YYYY-MM-DD). Defaults to last Monday.
    - week_end: End date (YYYY-MM-DD). Defaults to last Sunday.
    - salesperson_filter: Filter by salesperson name (optional)
    """
    try:
        config = Config()
        analyzer = WeeklyPipelineAnalyzer(config)

        report = analyzer.generate_weekly_report(
            week_start=week_start,
            week_end=week_end,
            salesperson_filter=salesperson_filter
        )

        return report

    except Exception as e:
        logger.error(f"Error generating weekly pipeline report: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/pipeline/send-to-teams")
def send_pipeline_to_teams(
    request: Dict[str, Any],
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """
    Send weekly pipeline report to Teams.

    Request body:
    {
        "chat_id": "19:xxx@thread.v2",
        "report_data": {...}
    }
    """
    try:
        chat_id = request.get('chat_id')
        report_data = request.get('report_data')

        if not chat_id or not report_data:
            raise HTTPException(status_code=400, detail="chat_id and report_data are required")

        # Get user's Microsoft access token
        token_store = EmailTokenStore()
        user_id = str(current_user.get("id"))

        tokens = token_store.get_tokens(user_id)

        if not tokens or not isinstance(tokens, dict):
            raise HTTPException(
                status_code=401,
                detail="No Microsoft authentication found. Please connect your Microsoft account in Settings."
            )

        access_token = tokens.get("access_token")

        if not access_token:
            raise HTTPException(
                status_code=401,
                detail="Microsoft token is invalid. Please reconnect your Microsoft account in Settings."
            )

        # Initialize Teams messenger
        teams = TeamsMessenger(access_token)

        # Format the report as HTML message
        message_html = TeamsMessenger.format_weekly_pipeline_report(report_data)

        # Send to Teams
        result = teams.send_message_to_chat(chat_id, message_html)

        logger.info(f"Successfully sent weekly pipeline report to Teams chat {chat_id}")

        return {
            "success": True,
            "message": "Pipeline report sent to Teams successfully",
            "teams_message_id": result.get("id")
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error sending pipeline report to Teams: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# NDA Analysis Endpoints
# ============================================================================

@app.post("/nda/upload")
async def upload_nda(
    file: UploadFile = File(...),
    save_to_database: bool = Form(True),
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: SupabaseDatabase = Depends(get_supabase_database)
):
    """Upload an NDA document for analysis. Optionally save to database."""
    try:
        from modules.nda_analyzer import NDAAnalyzer

        # Read file content
        file_content = await file.read()
        file_size = len(file_content)

        # Validate file size (max 10MB)
        if file_size > 10 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="File size exceeds 10MB limit")

        # Extract text from file
        analyzer = NDAAnalyzer()
        nda_text = analyzer.extract_text_from_file(file_content, file.filename)

        if not nda_text or len(nda_text.strip()) < 100:
            raise HTTPException(status_code=400, detail="Could not extract sufficient text from file. Please ensure the file contains readable text.")

        # Detect language
        language = analyzer.detect_language(nda_text)

        # Perform analysis
        try:
            analysis = analyzer.analyze_nda(nda_text, language)

            # If save_to_database is True, save the NDA and analysis
            if save_to_database:
                # Create NDA document record
                user_id = str(current_user.get("id", "unknown"))
                nda_id = db.create_nda_document(
                    user_id=user_id,
                    file_name=file.filename,
                    file_size=file_size,
                    file_content=nda_text,
                    language=language
                )

                if not nda_id:
                    raise HTTPException(status_code=500, detail="Failed to create NDA document record")

                # Save analysis results
                success = db.update_nda_analysis(
                    nda_id=nda_id,
                    risk_category=analysis["risk_category"],
                    risk_score=analysis["risk_score"],
                    summary=analysis["summary"],
                    questionable_clauses=analysis["questionable_clauses"],
                    analysis_details=analysis.get("raw_analysis", {}),
                    language=language
                )

                if not success:
                    raise Exception("Failed to save analysis results")

                # Get the complete document with analysis
                nda_doc = db.get_nda_document(nda_id)

                return {
                    "success": True,
                    "message": "NDA analyzed and saved successfully",
                    "nda": nda_doc,
                    "saved": True
                }
            else:
                # Return analysis without saving
                return {
                    "success": True,
                    "message": "NDA analyzed successfully (not saved)",
                    "nda": {
                        "file_name": file.filename,
                        "file_size": file_size,
                        "language": language,
                        "risk_category": analysis["risk_category"],
                        "risk_score": analysis["risk_score"],
                        "summary": analysis["summary"],
                        "questionable_clauses": analysis["questionable_clauses"],
                        "status": "completed"
                    },
                    "saved": False
                }

        except Exception as analysis_error:
            logger.error(f"Error analyzing NDA: {analysis_error}")
            if save_to_database and 'nda_id' in locals():
                db.update_nda_status(nda_id, "failed", str(analysis_error))
            raise HTTPException(status_code=500, detail=f"Analysis failed: {str(analysis_error)}")

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading NDA: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/nda/documents")
def get_nda_documents(
    limit: int = 50,
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: SupabaseDatabase = Depends(get_supabase_database)
):
    """Get list of NDA documents for the current user."""
    try:
        user_id = str(current_user.get("id", "unknown"))
        documents = db.get_nda_documents(user_id, limit)

        return {
            "success": True,
            "documents": documents,
            "count": len(documents)
        }
    except Exception as e:
        logger.error(f"Error getting NDA documents: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/nda/documents/{nda_id}")
def get_nda_document(
    nda_id: str,
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: SupabaseDatabase = Depends(get_supabase_database)
):
    """Get a specific NDA document by ID."""
    try:
        document = db.get_nda_document(nda_id)

        if not document:
            raise HTTPException(status_code=404, detail="NDA document not found")

        # Verify user owns this document
        user_id = str(current_user.get("id", "unknown"))
        if document.get("user_id") != user_id:
            raise HTTPException(status_code=403, detail="Access denied")

        return {
            "success": True,
            "document": document
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting NDA document: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/nda/documents/{nda_id}")
def delete_nda_document(
    nda_id: str,
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: SupabaseDatabase = Depends(get_supabase_database)
):
    """Delete an NDA document."""
    try:
        # Get document first to verify ownership
        document = db.get_nda_document(nda_id)

        if not document:
            raise HTTPException(status_code=404, detail="NDA document not found")

        # Verify user owns this document
        user_id = str(current_user.get("id", "unknown"))
        if document.get("user_id") != user_id:
            raise HTTPException(status_code=403, detail="Access denied")

        # Delete document
        success = db.delete_nda_document(nda_id)

        if not success:
            raise HTTPException(status_code=500, detail="Failed to delete document")

        return {
            "success": True,
            "message": "NDA document deleted successfully"
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting NDA document: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/proposal-followups/generate-draft")
def generate_email_draft(request: GenerateDraftRequest):
    """Generate an email draft for a specific thread."""
    try:
        analyzer = ProposalFollowupAnalyzer()
        analysis = analyzer.analyze_thread_with_llm(request.thread_data)

        return {
            "success": True,
            "draft": analysis.get("draft_email", ""),
            "analysis": analysis
        }

    except Exception as e:
        logger.error(f"Error generating draft: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/proposal-followups/refine-draft")
def refine_email_draft(request: RefineDraftRequest):
    """Refine an email draft based on user's editing instructions."""
    try:
        from modules.llm_client import LLMClient

        llm = LLMClient()

        # Create prompt for refinement
        prompt = f"""You are an expert email writer. Modify the following email draft based on the user's instructions.

Current Draft:
{request.current_draft}

User's Editing Instructions:
{request.edit_prompt}

Thread Context:
{json.dumps(request.thread_data, indent=2)}

Please provide the refined email draft that incorporates the user's requested changes while maintaining professionalism and the original intent."""

        messages = [
            {"role": "system", "content": "You are an expert email writer helping refine business communications."},
            {"role": "user", "content": prompt}
        ]

        refined_draft = llm.chat_completion(messages, max_tokens=1000, temperature=0.7)

        return {
            "success": True,
            "refined_draft": refined_draft
        }

    except Exception as e:
        logger.error(f"Error refining draft: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/proposal-followups/send-email")
def send_followup_email(
    request: SendFollowupEmailRequest,
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: SupabaseDatabase = Depends(get_supabase_database)
):
    """Send a follow-up email via Outlook and mark as completed."""
    try:
        from modules.outlook_client import OutlookClient
        from modules.token_store import TokenStore

        user_identifier = current_user.get("email")
        user_id = current_user.get("id")

        # Get user's Outlook token
        config = Config()
        token_store = TokenStore(config)
        tokens = token_store.get_tokens(user_identifier)

        if not tokens:
            raise HTTPException(
                status_code=401,
                detail="No Outlook authentication found. Please authenticate first."
            )

        outlook = OutlookClient(config)

        # Send the email as a reply
        result = outlook.send_reply(
            access_token=tokens.get("access_token"),
            conversation_id=request.conversation_id,
            reply_body=request.draft_body,
            subject=request.subject,
            reply_to_message_id=request.reply_to_message_id
        )

        if result:
            # Mark as completed
            db.mark_followup_complete(
                thread_id=request.conversation_id,
                conversation_id=request.conversation_id,
                user_id=user_id,
                completion_method="tool_sent",
                notes="Email sent via follow-up hub"
            )

            return {"success": True, "message": "Email sent successfully"}
        else:
            raise HTTPException(status_code=500, detail="Failed to send email")

    except Exception as e:
        logger.error(f"Error sending email: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# LEAD ASSIGNMENT ENDPOINTS
# ============================================================================

class LeadAssignmentCreate(BaseModel):
    conversation_id: str
    external_email: str
    subject: str
    assigned_to_user_id: Optional[int] = None  # Optional for Teams users
    assigned_to_teams_id: Optional[str] = None  # Azure AD user ID for Teams
    assigned_to_name: Optional[str] = None  # Display name
    assigned_to_email: Optional[str] = None  # Email address
    lead_data: Dict[str, Any]
    notes: Optional[str] = None
    analysis_cache_id: Optional[str] = None


class LeadAssignmentUpdate(BaseModel):
    status: Literal["accepted", "completed", "rejected"]
    notes: Optional[str] = None


@app.post("/lead-assignments")
def create_lead_assignment(
    assignment: LeadAssignmentCreate,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """
    Assign a lead to another user.

    Args:
        assignment: Assignment details
        current_user: Authenticated user making the assignment

    Returns:
        Created assignment record
    """
    if not supabase.is_connected():
        raise HTTPException(
            status_code=503,
            detail="Lead assignment requires Supabase connection"
        )

    try:
        user_id = current_user.get("id")

        # Build assignment data with Teams support
        assignment_kwargs = {
            "conversation_id": assignment.conversation_id,
            "external_email": assignment.external_email,
            "subject": assignment.subject,
            "assigned_from_user_id": user_id,
            "lead_data": assignment.lead_data,
            "notes": assignment.notes,
            "analysis_cache_id": assignment.analysis_cache_id
        }

        # Add user ID or Teams info
        if assignment.assigned_to_user_id:
            assignment_kwargs["assigned_to_user_id"] = assignment.assigned_to_user_id
        elif assignment.assigned_to_teams_id:
            assignment_kwargs["assigned_to_teams_id"] = assignment.assigned_to_teams_id
            assignment_kwargs["assigned_to_name"] = assignment.assigned_to_name
            assignment_kwargs["assigned_to_email"] = assignment.assigned_to_email
        else:
            raise HTTPException(status_code=400, detail="Must provide either assigned_to_user_id or assigned_to_teams_id")

        result = supabase.create_lead_assignment(**assignment_kwargs)

        if not result:
            raise HTTPException(status_code=500, detail="Failed to create assignment")

        return result

    except Exception as e:
        logger.error(f"Error creating lead assignment: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/lead-assignments/received")
def get_received_assignments(
    status: Optional[Literal["pending", "accepted", "completed", "rejected"]] = None,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """
    Get leads assigned to the current user.

    Args:
        status: Optional filter by status
        current_user: Authenticated user

    Returns:
        List of received assignments
    """
    if not supabase.is_connected():
        raise HTTPException(
            status_code=503,
            detail="Lead assignment requires Supabase connection"
        )

    try:
        user_id = current_user.get("id")
        assignments = supabase.get_received_assignments(user_id=user_id, status=status)
        return {"assignments": assignments, "count": len(assignments)}

    except Exception as e:
        logger.error(f"Error fetching received assignments: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/lead-assignments/sent")
def get_sent_assignments(
    status: Optional[Literal["pending", "accepted", "completed", "rejected"]] = None,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """
    Get leads assigned by the current user.

    Args:
        status: Optional filter by status
        current_user: Authenticated user

    Returns:
        List of sent assignments
    """
    if not supabase.is_connected():
        raise HTTPException(
            status_code=503,
            detail="Lead assignment requires Supabase connection"
        )

    try:
        user_id = current_user.get("id")
        assignments = supabase.get_sent_assignments(user_id=user_id, status=status)
        return {"assignments": assignments, "count": len(assignments)}

    except Exception as e:
        logger.error(f"Error fetching sent assignments: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@app.patch("/lead-assignments/{assignment_id}")
def update_assignment(
    assignment_id: str,
    update: LeadAssignmentUpdate,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """
    Update a lead assignment status.

    Args:
        assignment_id: Assignment UUID
        update: Status update
        current_user: Authenticated user

    Returns:
        Success status
    """
    if not supabase.is_connected():
        raise HTTPException(
            status_code=503,
            detail="Lead assignment requires Supabase connection"
        )

    try:
        success = supabase.update_assignment_status(
            assignment_id=assignment_id,
            status=update.status,
            notes=update.notes
        )

        if not success:
            raise HTTPException(status_code=500, detail="Failed to update assignment")

        return {"success": True, "assignment_id": assignment_id, "status": update.status}

    except Exception as e:
        logger.error(f"Error updating assignment: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# EMAIL / OUTLOOK OAUTH2 ENDPOINTS
# ============================================================================

def get_outlook_client() -> OutlookClient:
    """Get or create Outlook client instance."""
    return OutlookClient(config=Config())


def get_token_store() -> EmailTokenStore:
    """Get or create token store instance."""
    return EmailTokenStore()


class EmailAuthResponse(BaseModel):
    authorization_url: str
    state: str


class EmailAuthCallbackRequest(BaseModel):
    code: str
    state: str
    user_identifier: Optional[str] = None  # e.g., Odoo user email or ID


class EmailAuthStatusResponse(BaseModel):
    authorized: bool
    user_email: Optional[str] = None
    user_name: Optional[str] = None
    expires_soon: bool = False


@app.get("/auth/outlook/start", response_model=EmailAuthResponse)
def start_outlook_auth(current_user: Dict[str, Any] = Depends(get_current_user)):
    """
    Start OAuth2 flow for Outlook/Microsoft email.
    Returns authorization URL for user to visit.
    Requires authentication.
    """
    import secrets

    outlook = get_outlook_client()

    # Generate random state for CSRF protection
    # Include user ID in state for verification
    state = f"{current_user['id']}:{secrets.token_urlsafe(32)}"

    auth_url = outlook.get_authorization_url(state=state)

    return EmailAuthResponse(
        authorization_url=auth_url,
        state=state
    )


@app.get("/auth/outlook/callback", response_class=RedirectResponse)
async def outlook_auth_callback(request: Request, code: str, state: str):
    """
    Handle OAuth2 callback from Microsoft (GET with query params).
    Redirects to frontend callback page which will complete the flow.
    """
    # Redirect to frontend with the code and state
    # Frontend will call the POST endpoint to complete the exchange
    cfg = Config()
    frontend_url = cfg.FRONTEND_URL.rstrip('/')
    frontend_callback_url = f"{frontend_url}/auth/outlook/callback?code={code}&state={state}"
    return RedirectResponse(url=frontend_callback_url)


@app.post("/auth/outlook/callback")
def outlook_auth_callback_post(
    request: EmailAuthCallbackRequest,
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: Database = Depends(get_database)
):
    """
    Complete OAuth2 callback (POST from frontend).
    Exchange authorization code for tokens and store them for the current user.
    Requires authentication.
    """
    outlook = get_outlook_client()
    token_store = get_token_store()

    try:
        # Verify state contains current user ID
        if request.state and ":" in request.state:
            state_user_id = int(request.state.split(":")[0])
            if state_user_id != current_user["id"]:
                raise HTTPException(status_code=403, detail="State verification failed")

        # Exchange code for tokens
        token_response = outlook.exchange_code_for_tokens(request.code)

        access_token = token_response.get("access_token")
        refresh_token = token_response.get("refresh_token")
        expires_in = token_response.get("expires_in", 3600)

        if not access_token or not refresh_token:
            raise HTTPException(status_code=400, detail="Invalid token response from Microsoft")

        # Get user info from Microsoft
        user_info = outlook.get_user_info(access_token)
        user_email = user_info.get("mail") or user_info.get("userPrincipalName")
        user_name = user_info.get("displayName")

        # Use user ID as identifier for token storage
        user_identifier = str(current_user["id"])

        # Store tokens in file system (backwards compatibility)
        success = token_store.save_tokens(
            user_identifier=user_identifier,
            access_token=access_token,
            refresh_token=refresh_token,
            expires_in=expires_in,
            user_email=user_email,
            user_name=user_name,
        )

        if not success:
            raise HTTPException(status_code=500, detail="Failed to store tokens")

        # Also store in database (with expires_at for persistence)
        from datetime import datetime, timedelta
        expires_at = (datetime.utcnow() + timedelta(seconds=expires_in)).isoformat()

        try:
            db.update_user_settings(
                user_id=current_user["id"],
                outlook_tokens={
                    "access_token": access_token,
                    "refresh_token": refresh_token,
                    "expires_in": expires_in,
                    "expires_at": expires_at,
                    "user_email": user_email,
                    "user_name": user_name,
                    "created_at": datetime.utcnow().isoformat(),
                    "updated_at": datetime.utcnow().isoformat()
                },
                user_identifier=user_identifier
            )
            logger.info(f"✅ Successfully saved Outlook tokens to database for user {current_user['id']} ({user_email})")
        except Exception as db_error:
            logger.error(f"❌ Failed to save Outlook tokens to database: {db_error}")
            # Don't fail the request if database save fails, file system backup exists

        return {
            "success": True,
            "user_email": user_email,
            "user_name": user_name,
            "message": "Email authorization successful"
        }

    except Exception as e:
        logger.error(f"Error in OAuth callback: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/auth/outlook/status", response_model=EmailAuthStatusResponse)
def get_email_auth_status(current_user: Dict[str, Any] = Depends(get_current_user)):
    """Check if current user has authorized email access."""
    token_store = get_token_store()
    user_identifier = str(current_user["id"])

    tokens = token_store.get_tokens(user_identifier)

    if not tokens:
        return EmailAuthStatusResponse(authorized=False)

    is_expired = token_store.is_token_expired(user_identifier)

    return EmailAuthStatusResponse(
        authorized=True,
        user_email=tokens.get("user_email"),
        user_name=tokens.get("user_name"),
        expires_soon=is_expired
    )


@app.delete("/auth/outlook")
def revoke_email_auth(
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: Database = Depends(get_database)
):
    """Revoke email authorization for current user."""
    token_store = get_token_store()
    user_identifier = str(current_user["id"])

    success = token_store.delete_tokens(user_identifier)

    if not success:
        raise HTTPException(status_code=500, detail="Failed to revoke authorization")

    # Also clear from database
    db.update_user_settings(
        user_id=current_user["id"],
        outlook_tokens=None
    )

    return {"success": True, "message": "Email authorization revoked"}


# ============================================================================
# SYSTEM EMAIL OAUTH2 ENDPOINTS (for automated.response@prezlab.com)
# ============================================================================

@app.get("/auth/outlook/system/start", response_model=EmailAuthResponse)
def start_system_outlook_auth(current_user: Dict[str, Any] = Depends(get_current_user)):
    """
    Start OAuth2 flow for system email (automated.response@prezlab.com).
    Admin only. Returns authorization URL for user to visit.
    """
    # Could add admin check here if needed
    # if current_user["role"] != "admin":
    #     raise HTTPException(status_code=403, detail="Admin access required")

    import secrets

    outlook = get_outlook_client()

    # Use special state prefix to identify system auth
    state = f"SYSTEM:{secrets.token_urlsafe(32)}"

    # Force account selection and exclude Teams permissions for system email
    # System email only needs Mail.Read for engage group inbox access
    auth_url = outlook.get_authorization_url(
        state=state,
        force_account_selection=True,
        include_teams=False  # No Teams permissions needed for system email
    )

    return EmailAuthResponse(
        authorization_url=auth_url,
        state=state
    )


@app.post("/auth/outlook/system/callback")
def system_outlook_auth_callback_post(
    request: EmailAuthCallbackRequest,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """
    Complete OAuth2 callback for system email (POST from frontend).
    Exchange authorization code for tokens and store them in Supabase for all users.
    """
    outlook = get_outlook_client()

    try:
        # Verify state is for system auth
        if not request.state or not request.state.startswith("SYSTEM:"):
            raise HTTPException(status_code=403, detail="Invalid state for system auth")

        # Exchange code for tokens
        token_response = outlook.exchange_code_for_tokens(request.code)

        access_token = token_response.get("access_token")
        refresh_token = token_response.get("refresh_token")
        expires_in = token_response.get("expires_in", 3600)

        if not access_token or not refresh_token:
            raise HTTPException(status_code=400, detail="Invalid token response from Microsoft")

        # Get user info from Microsoft
        user_info = outlook.get_user_info(access_token)
        user_email = user_info.get("mail") or user_info.get("userPrincipalName")
        user_name = user_info.get("displayName")

        # Use fixed identifier for system account
        system_identifier = "SYSTEM_automated.response@prezlab.com"

        # Store tokens in Supabase (persists across deployments)
        from datetime import datetime, timedelta, timezone
        expires_at = datetime.now(timezone.utc) + timedelta(seconds=expires_in)

        try:
            # Upsert into email_tokens table
            result = supabase.client.table("email_tokens").upsert({
                "user_identifier": system_identifier,
                "access_token": access_token,
                "refresh_token": refresh_token,
                "expires_at": expires_at.isoformat(),
                "user_email": user_email,
                "user_name": user_name,
                "updated_at": datetime.now(timezone.utc).isoformat()
            }, on_conflict="user_identifier").execute()

            logger.info(f"System email tokens stored successfully for {user_email}")
        except Exception as e:
            logger.error(f"Failed to store system email tokens: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to store tokens: {str(e)}")

        return {
            "success": True,
            "user_email": user_email,
            "user_name": user_name,
            "message": "System email authorization successful"
        }

    except Exception as e:
        logger.error(f"Error in system OAuth callback: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/auth/outlook/system/status", response_model=EmailAuthStatusResponse)
def get_system_email_auth_status():
    """Check if system email (automated.response@prezlab.com) has authorized access."""
    outlook = get_outlook_client()
    system_identifier = "SYSTEM_automated.response@prezlab.com"

    tokens = outlook.get_user_auth_tokens(system_identifier)

    if not tokens:
        return EmailAuthStatusResponse(authorized=False)

    # Check if token is expired (will be refreshed automatically on next use)
    expires_at = tokens.get("expires_at")
    expires_soon = False
    if expires_at:
        from datetime import datetime, timezone, timedelta
        expires_dt = datetime.fromisoformat(expires_at.replace('Z', '+00:00'))
        now = datetime.now(timezone.utc)
        expires_soon = (expires_dt - now) < timedelta(minutes=10)

    return EmailAuthStatusResponse(
        authorized=True,
        user_email=tokens.get("user_email"),
        user_name=tokens.get("user_name"),
        expires_soon=expires_soon
    )


@app.delete("/auth/outlook/system")
def revoke_system_email_auth(current_user: Dict[str, Any] = Depends(get_current_user)):
    """Revoke system email authorization. Admin only."""
    outlook = get_outlook_client()
    system_identifier = "SYSTEM_automated.response@prezlab.com"

    success = outlook.delete_user_auth_tokens(system_identifier)

    if not success:
        raise HTTPException(status_code=500, detail="Failed to revoke system authorization")

    return {"success": True, "message": "System email authorization revoked"}


@app.get("/auth/outlook/users")
def list_authorized_users():
    """List all users who have authorized email access (admin endpoint)."""
    token_store = get_token_store()
    users = token_store.list_authorized_users()
    return {"users": users, "count": len(users)}


class SendEmailRequest(BaseModel):
    to: List[str]
    subject: str
    body: str
    cc: Optional[List[str]] = None
    bcc: Optional[List[str]] = None
    include_engage_cc: bool = True  # Automatically CC engage@prezlab.com


@app.post("/email/send")
def send_email(
    request: SendEmailRequest,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """
    Send an email from current user with automatic CC to engage@prezlab.com.
    Requires authentication and email authorization.
    """
    outlook = get_outlook_client()
    token_store = get_token_store()
    user_identifier = str(current_user["id"])

    try:
        # Get tokens
        tokens = token_store.get_tokens(user_identifier)
        if not tokens:
            raise HTTPException(
                status_code=400,
                detail="Email not authorized. Please connect your email first."
            )

        # Refresh token if expired
        access_token = tokens.get("access_token")
        if token_store.is_token_expired(user_identifier):
            refresh_token = tokens.get("refresh_token")
            token_response = outlook.refresh_access_token(refresh_token)
            access_token = token_response.get("access_token")
            token_store.update_access_token(
                user_identifier,
                access_token,
                token_response.get("expires_in", 3600)
            )

        # Add engage@prezlab.com to CC if requested
        cc_list = list(request.cc or [])
        if request.include_engage_cc and "engage@prezlab.com" not in cc_list:
            cc_list.append("engage@prezlab.com")

        # Send email
        success = outlook.send_email(
            access_token=access_token,
            to=request.to,
            subject=request.subject,
            body=request.body,
            cc=cc_list if cc_list else None,
            bcc=request.bcc
        )

        if not success:
            raise HTTPException(status_code=500, detail="Failed to send email")

        return {
            "success": True,
            "message": "Email sent successfully",
            "cc_list": cc_list
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error sending email: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# CALL FLOW GENERATION ENDPOINTS
# ============================================================================

class EnrichedLead(BaseModel):
    id: int
    name: str
    partner_name: Optional[str] = None
    email_from: Optional[str] = None
    stage_name: Optional[str] = None
    salesperson_name: Optional[str] = None
    description: Optional[str] = None
    function: Optional[str] = None


class EnrichedLeadsResponse(BaseModel):
    count: int
    leads: List[EnrichedLead]


class CallFlowGenerateRequest(BaseModel):
    lead_id: int


@app.get("/enriched-leads", response_model=EnrichedLeadsResponse)
def get_enriched_leads():
    """Get list of leads that have been enriched (have quality rating)."""
    config = _setup_logging()
    odoo = OdooClient(config)

    if not odoo.connect():
        raise HTTPException(status_code=500, detail="Failed to connect to Odoo")

    try:
        # Find Dareen's user ID
        salesperson_name = config.SALESPERSON_NAME
        user_id = odoo.find_user_id(salesperson_name)

        # Enriched leads are those that either:
        # 1. Have a quality rating (x_studio_quality is not empty), OR
        # 2. Don't belong to Dareen
        # This is the opposite of unenriched leads
        domain = [
            '|',
            ['x_studio_quality', '!=', False],
            ['user_id', '!=', user_id] if user_id else ['id', '!=', -1]
        ]

        fields = [
            'id', 'name', 'partner_name', 'email_from', 'stage_id',
            'user_id', 'description', 'function', 'x_studio_quality'
        ]

        leads_data = odoo._call_kw(
            'crm.lead',
            'search_read',
            [domain],
            {
                'fields': fields,
                'limit': 100,
                'order': 'write_date desc'
            }
        )

        enriched_leads = []
        for lead in leads_data:
            enriched_leads.append(
                EnrichedLead(
                    id=lead.get('id'),
                    name=lead.get('name') or 'Untitled Lead',
                    partner_name=lead.get('partner_name') if isinstance(lead.get('partner_name'), str) else None,
                    email_from=lead.get('email_from') if isinstance(lead.get('email_from'), str) else None,
                    stage_name=lead.get('stage_id')[1] if lead.get('stage_id') and isinstance(lead.get('stage_id'), (list, tuple)) else None,
                    salesperson_name=lead.get('user_id')[1] if lead.get('user_id') and isinstance(lead.get('user_id'), (list, tuple)) else None,
                    description=lead.get('description') if isinstance(lead.get('description'), str) else None,
                    function=lead.get('function') if isinstance(lead.get('function'), str) else None,
                )
            )

        return EnrichedLeadsResponse(count=len(enriched_leads), leads=enriched_leads)

    except Exception as exc:
        logger.error(f"Failed to fetch enriched leads: {exc}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch enriched leads: {str(exc)}")


@app.post("/call-flow/generate")
def generate_call_flow(request: CallFlowGenerateRequest):
    """Generate a personalized discovery call flow document for a lead."""
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from openai import OpenAI

    config = _setup_logging()
    odoo = OdooClient(config)

    if not odoo.connect():
        raise HTTPException(status_code=500, detail="Failed to connect to Odoo")

    # Fetch the lead data
    try:
        lead_data = odoo._call_kw(
            'crm.lead',
            'search_read',
            [[('id', '=', request.lead_id)]],
            {
                'fields': ['id', 'name', 'partner_name', 'email_from', 'stage_id',
                          'user_id', 'description', 'function', 'phone', 'mobile'],
                'limit': 1
            }
        )

        if not lead_data:
            raise HTTPException(status_code=404, detail="Lead not found")

        lead = lead_data[0]

    except Exception as exc:
        logger.error(f"Failed to fetch lead {request.lead_id}: {exc}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch lead: {str(exc)}")

    # Extract lead information
    lead_name = lead.get('name') or 'Unknown'
    partner_name = lead.get('partner_name') if isinstance(lead.get('partner_name'), str) else 'Unknown Company'
    description = lead.get('description') if isinstance(lead.get('description'), str) else ''
    job_title = lead.get('function') if isinstance(lead.get('function'), str) else 'Unknown'
    stage = lead.get('stage_id')[1] if lead.get('stage_id') and isinstance(lead.get('stage_id'), (list, tuple)) else 'Unknown'

    # Get knowledge base context about PrezLab
    kb_context = ""
    try:
        supabase = SupabaseClient()
        if supabase.is_connected():
            result = supabase.client.table("knowledge_base_documents")\
                .select("filename, content")\
                .eq("is_active", True)\
                .execute()

            context_parts = []
            for doc in result.data:
                filename = doc.get("filename", "Unknown Document")
                content = doc.get("content", "")
                if content.strip():
                    context_parts.append(f"=== {filename} ===\n{content.strip()}")

            kb_context = "\n\n".join(context_parts)
    except Exception as e:
        logger.warning(f"Could not fetch knowledge base context: {e}")

    # Generate personalized content using LLM
    try:
        openai_client = OpenAI(api_key=config.OPENAI_API_KEY)

        prompt_parts = []

        if kb_context:
            prompt_parts.append(f"""IMPORTANT - Company Context & Knowledge Base:
First, carefully review the following documents about PrezLab (our company, services, and offerings):

{kb_context}

INSTRUCTIONS:
- Study the above documents to understand what PrezLab offers and our value propositions
- Tailor your discovery questions to uncover needs that align with PrezLab's specific services
- Reference PrezLab's actual capabilities when framing questions about solutions
- Ensure all questions are relevant to what PrezLab can actually provide

""")

        prompt_parts.append(f"""You are a PrezLab sales consultant preparing for a discovery call. Using a consultative, dialogue-friendly approach, create personalized questions for each section that help uncover the prospect's business challenges and how PrezLab's services can help.

LEAD INFORMATION:
- Contact Name: {lead_name}
- Company: {partner_name}
- Job Title: {job_title}
- Current Stage: {stage}
- Enriched Notes: {description[:500] if description else 'No additional context available'}

DISCOVERY CALL FLOW STRUCTURE (based on proven methodology):

1. **Business Problem** - What challenge or opportunity is driving them to explore working with us?
   - Focus on understanding their core pain point or strategic need
   - Make it conversational and open-ended
   - Examples: challenges with communication clarity, brand consistency, executive presentations, etc.

2. **Current State** - How are they handling things today?
   - Understand their current approach (in-house, agencies, mix)
   - What's working and what could be stronger

3. **Cause Analysis** - What's the root cause?
   - Help them identify if it's about resources, processes, alignment, or capabilities
   - Be empathetic and avoid making them defensive

4. **Negative Impact** - What happens when these challenges occur?
   - Explore consequences: delays, rework, missed opportunities, brand inconsistency
   - Help them articulate the cost of inaction

5. **Desired Outcome** - What does success look like?
   - Future state vision in 6 months
   - What outcomes would make this "worth the investment"

6. **Process** - How do they like to work with partners?
   - Project-based vs. ongoing engagement preferences
   - Timeline, review, and approval processes

7. **Stakeholders** - Who else is involved?
   - Decision-makers, influencers, departments affected
   - Identify champions and potential blockers

For each section, provide:
- A brief **objective** (what you're trying to discover)
- 2-3 **dialogue-friendly questions** that are:
  * Conversational and consultative (not interrogative)
  * Tailored to {job_title} at {partner_name}
  * Relevant to PrezLab's services (presentations, visual communication, content)
  * Open-ended to encourage discussion

Format as JSON:
{{
  "sections": [
    {{
      "title": "Business Problem",
      "objective": "...",
      "questions": ["...", "..."]
    }},
    ...
  ]
}}

Make questions sound natural in conversation. Reference their context from enriched notes when relevant. Focus on business impact, not just creative/design needs.""")

        prompt = "\n".join(prompt_parts)

        response = openai_client.chat.completions.create(
            model=config.OPENAI_MODEL or 'gpt-5-mini',
            messages=[
                {"role": "system", "content": "You are an expert sales consultant who creates personalized discovery call frameworks. Always respond with valid JSON."},
                {"role": "user", "content": prompt}
            ],
            max_completion_tokens=2000,
            response_format={"type": "json_object"}
        )

        import json
        call_flow_data = json.loads(response.choices[0].message.content)

    except Exception as exc:
        logger.error(f"Failed to generate call flow content: {exc}")
        raise HTTPException(status_code=500, detail=f"Failed to generate call flow: {str(exc)}")

    # Create the Word document
    try:
        doc = Document()

        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Title
        title = doc.add_heading(f'Discovery Call Flow - {lead_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle with company info
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f'{partner_name} | {job_title}')
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()  # Spacing

        # Add each section
        sections_data = call_flow_data.get('sections', [])
        for i, section_data in enumerate(sections_data, 1):
            # Section title
            section_title = doc.add_heading(f"{i}. {section_data.get('title', '')}", 1)
            section_title.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Objective
            objective_para = doc.add_paragraph()
            objective_para.add_run('Objective: ').bold = True
            objective_para.add_run(section_data.get('objective', ''))

            # Questions
            questions_heading = doc.add_paragraph()
            questions_heading.add_run('Discussion Questions:').bold = True

            questions = section_data.get('questions', [])
            for question in questions:
                doc.add_paragraph(question, style='List Bullet')

            doc.add_paragraph()  # Spacing between sections

        # Add footer with preparation notes
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run('Prepared by PrezLab Lead Automation Hub')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(150, 150, 150)

        # Save to bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Return as downloadable file
        filename = f"Discovery_Call_Flow_{lead_name.replace(' ', '_')}.docx"

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    except Exception as exc:
        logger.error(f"Failed to create document: {exc}")
        raise HTTPException(status_code=500, detail=f"Failed to create document: {str(exc)}")


# ============================================================================
# DASHBOARD SUMMARY ENDPOINT
# ============================================================================

class DashboardSummary(BaseModel):
    """Dashboard summary data aggregated from all features."""
    high_priority_count: int
    high_priority_items: List[Dict[str, Any]]
    stats: Dict[str, Any]  # Now includes last_updated timestamp
    recent_activity: List[Dict[str, Any]]


@app.get("/dashboard/summary", response_model=DashboardSummary)
def get_dashboard_summary(
    engage_email: str = "automated.response@prezlab.com"
):
    """
    Get dashboard summary with data from all features.

    Args:
        engage_email: Email address for engage group monitoring

    Returns:
        Aggregated dashboard data
    """
    try:
        high_priority_items = []
        stats = {}
        recent_activity = []

        # 1. Get Proposal Follow-ups data from Supabase (latest shared report)
        try:
            # Try to get the latest shared report from Supabase
            if supabase.is_connected():
                latest_report_response = supabase.client.table("analysis_cache") \
                    .select("*") \
                    .eq("analysis_type", "proposal_followups") \
                    .eq("is_shared", True) \
                    .not_.is_("report_type", "null") \
                    .order("created_at", desc=True) \
                    .limit(1) \
                    .execute()
            else:
                latest_report_response = None

            if latest_report_response and latest_report_response.data and len(latest_report_response.data) > 0:
                report_data = latest_report_response.data[0]

                # The field is 'results' and it's JSON-encoded
                import json
                result = report_data.get("results")
                if isinstance(result, str):
                    result = json.loads(result)
                if not result:
                    result = {}

                unanswered_count = result.get("summary", {}).get("unanswered_count", 0)
                pending_proposals_count = result.get("summary", {}).get("pending_proposals_count", 0)
                stats["unanswered_emails"] = unanswered_count
                stats["pending_proposals"] = pending_proposals_count
                stats["last_updated"] = report_data.get("created_at")

                followups_result = {
                    "summary": {
                        "unanswered_count": unanswered_count,
                        "pending_proposals_count": pending_proposals_count
                    },
                    "unanswered": result.get("unanswered", []),
                    "pending_proposals": result.get("pending_proposals", [])
                }
            else:
                # No reports yet
                stats["unanswered_emails"] = 0
                stats["pending_proposals"] = 0
                stats["last_updated"] = None
                followups_result = {"summary": {"unanswered_count": 0, "pending_proposals_count": 0}, "unanswered": [], "pending_proposals": []}

            # Add high priority items (>5 days waiting)
            for item in followups_result["unanswered"]:
                if item["days_waiting"] >= 5:
                    high_priority_items.append({
                        "type": "email",
                        "subject": item["subject"],
                        "external_email": item["external_email"],
                        "days_waiting": item["days_waiting"],
                        "odoo_lead": item.get("odoo_lead"),
                        "source": "engage"
                    })

            for item in followups_result["pending_proposals"]:
                if item["days_waiting"] >= 5:
                    high_priority_items.append({
                        "type": "proposal",
                        "subject": item["subject"],
                        "external_email": item["external_email"],
                        "days_waiting": item["days_waiting"],
                        "odoo_lead": item.get("odoo_lead"),
                        "source": "engage"
                    })

            # Add recent activity
            for item in followups_result["unanswered"][:3]:  # Latest 3
                recent_activity.append({
                    "type": "email_received",
                    "description": f"Email from {item['external_email']}",
                    "time": item.get("last_contact_date", ""),
                    "subject": item["subject"]
                })

        except Exception as e:
            logger.error(f"Error fetching proposal follow-ups for dashboard: {e}")
            stats["unanswered_emails"] = 0
            stats["pending_proposals"] = 0

        # 2. Get Lost Leads data - fetch actual count from Odoo
        try:
            analyzer = get_lost_lead_analyzer()
            # Get lost leads without limit to count all
            lost_leads = analyzer.list_lost_leads(limit=1000)
            stats["lost_leads"] = len(lost_leads)
        except Exception as e:
            logger.error(f"Error fetching lost leads for dashboard: {e}")
            stats["lost_leads"] = 0

        # 3. Get Unenriched leads count (replacing enriched_today)
        try:
            workflow = get_workflow()
            _, unenriched_leads = workflow.generate_enrichment_prompt()
            stats["unenriched_leads"] = len(unenriched_leads)
        except Exception as e:
            logger.error(f"Error fetching unenriched leads for dashboard: {e}")
            stats["unenriched_leads"] = 0

        # 4. Get Call Flow stats (placeholder)
        try:
            # TODO: Track call flows in database and fetch count
            stats["call_flows_generated"] = 0
        except Exception as e:
            logger.error(f"Error fetching call flow stats for dashboard: {e}")
            stats["call_flows_generated"] = 0

        # Sort high priority items by days waiting (descending)
        high_priority_items.sort(key=lambda x: x["days_waiting"], reverse=True)

        # Sort recent activity by time (most recent first)
        recent_activity.sort(key=lambda x: x.get("time", ""), reverse=True)

        return DashboardSummary(
            high_priority_count=len(high_priority_items),
            high_priority_items=high_priority_items[:10],  # Limit to 10
            stats=stats,
            recent_activity=recent_activity[:10]  # Limit to 10
        )

    except Exception as e:
        logger.error(f"Error getting dashboard summary: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ===================================================================
# Authentication Endpoints
# ===================================================================

class LoginRequest(BaseModel):
    email: str
    password: str


class RegisterRequest(BaseModel):
    email: str
    name: str
    password: str


class LoginResponse(BaseModel):
    access_token: str
    token_type: str
    refresh_token: str
    user: Dict[str, Any]


class UserResponse(BaseModel):
    id: int
    email: str
    name: str
    role: str


@app.post("/auth/login", response_model=LoginResponse)
def login(request: LoginRequest, auth_service: AuthService = Depends(get_auth_service), db: Database = Depends(get_database)):
    """Authenticate user against Odoo and return access token."""
    # Try to authenticate against Odoo using XML-RPC (no session conflicts)
    from modules.odoo_client import OdooClient
    from config import Config
    import xmlrpc.client

    config = Config()
    base_url = config.ODOO_URL.rstrip('/')

    try:
        # Use XML-RPC for authentication - this won't create web sessions
        # that conflict with other applications
        common = xmlrpc.client.ServerProxy(f'{base_url}/xmlrpc/2/common', allow_none=True)

        # Authenticate and get UID
        odoo_uid = common.authenticate(
            config.ODOO_DB,
            request.email,
            request.password,
            {}
        )

        if not odoo_uid:
            raise HTTPException(status_code=401, detail="Invalid Odoo credentials")

        # Get user info
        models = xmlrpc.client.ServerProxy(f'{base_url}/xmlrpc/2/object', allow_none=True)
        user_info = models.execute_kw(
            config.ODOO_DB,
            odoo_uid,
            request.password,
            'res.users',
            'read',
            [odoo_uid],
            {'fields': ['name', 'email']}
        )

        odoo_name = user_info[0].get('name', request.email) if user_info else request.email

        if not odoo_uid:
            raise HTTPException(status_code=401, detail="Invalid Odoo credentials")

        # Check if user exists in our database
        user = auth_service.db.get_user_by_email(request.email)

        if not user:
            # Create new user automatically
            logger.info(f"Creating new user for Odoo account: {request.email}")
            password_hash = auth_service.hash_password(request.password)
            user_id = auth_service.db.create_user(
                email=request.email,
                name=odoo_name,
                password_hash=password_hash,
                role="user"
            )
            user = auth_service.db.get_user_by_id(user_id)

        # Store Odoo credentials in user settings
        db.update_user_settings(
            user_id=user["id"],
            odoo_url=config.ODOO_URL,
            odoo_db=config.ODOO_DB,
            odoo_username=request.email,
            odoo_password=request.password  # In production, encrypt this
        )

        # Update last login
        auth_service.db.update_last_login(user["id"])

        # Create access token
        access_token = auth_service.create_access_token(
            user_id=user["id"],
            email=user["email"],
            role=user["role"]
        )

        # Create refresh token for persistent login (never expires)
        refresh_token = auth_service.create_refresh_token()
        device_info = f"Teams App"  # Could be enhanced with more details
        auth_service.db.create_refresh_token(user["id"], refresh_token, device_info)

        return LoginResponse(
            access_token=access_token,
            token_type="bearer",
            refresh_token=refresh_token,
            user={
                "id": user["id"],
                "email": user["email"],
                "name": user["name"],
                "role": user["role"]
            }
        )

    except requests.RequestException as e:
        logger.error(f"Odoo authentication failed: {e}")
        raise HTTPException(status_code=401, detail="Unable to authenticate with Odoo")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Login error: {e}")
        raise HTTPException(status_code=500, detail="Authentication failed")


class RefreshTokenRequest(BaseModel):
    refresh_token: str


@app.post("/auth/refresh", response_model=LoginResponse)
def refresh_token(request: RefreshTokenRequest, auth_service: AuthService = Depends(get_auth_service)):
    """Exchange refresh token for new access token (persistent login)."""
    result = auth_service.refresh_access_token(request.refresh_token)

    if not result:
        raise HTTPException(status_code=401, detail="Invalid or expired refresh token")

    return LoginResponse(
        access_token=result["access_token"],
        token_type="bearer",
        refresh_token=request.refresh_token,  # Return same refresh token (it doesn't expire)
        user=result["user"]
    )


@app.post("/auth/register", response_model=UserResponse)
def register(
    request: RegisterRequest,
    auth_service: AuthService = Depends(get_auth_service)
):
    """Register a new user."""
    try:
        user_id = auth_service.register_user(
            email=request.email,
            name=request.name,
            password=request.password,
            role="user"
        )

        user = auth_service.db.get_user_by_id(user_id)
        return UserResponse(
            id=user["id"],
            email=user["email"],
            name=user["name"],
            role=user["role"]
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.get("/auth/me", response_model=UserResponse)
def get_me(current_user: Dict[str, Any] = Depends(get_current_user)):
    """Get current authenticated user."""
    return UserResponse(
        id=current_user["id"],
        email=current_user["email"],
        name=current_user["name"],
        role=current_user["role"]
    )


@app.get("/auth/users")
def list_users(
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: Database = Depends(get_database)
):
    """List all users (admin only)."""
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    return db.list_users()


# ============================================================================
# KNOWLEDGE BASE ENDPOINTS
# ============================================================================

class KnowledgeBaseDocument(BaseModel):
    """Knowledge base document metadata."""
    id: str
    filename: str
    file_size: int
    description: Optional[str] = None
    uploaded_by_user_id: int
    is_active: bool
    created_at: str
    updated_at: str


class KnowledgeBaseUploadResponse(BaseModel):
    """Response for document upload."""
    success: bool
    document_id: Optional[str] = None
    message: Optional[str] = None


@app.post("/knowledge-base/upload", response_model=KnowledgeBaseUploadResponse)
async def upload_knowledge_base_document(
    file: UploadFile = File(...),
    description: Optional[str] = Form(None),
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """
    Upload a PDF document to the knowledge base.
    The content will be extracted and used as context for AI analyses.
    """
    if not supabase.is_connected():
        raise HTTPException(
            status_code=503,
            detail="Knowledge base requires Supabase connection"
        )

    # Validate file type
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(
            status_code=400,
            detail="Only PDF files are supported"
        )

    try:
        # Read file content
        content_bytes = await file.read()
        file_size = len(content_bytes)

        # Extract text from PDF
        import io
        try:
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content_bytes))
            text_content = ""
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"

            if not text_content.strip():
                raise HTTPException(
                    status_code=400,
                    detail="Could not extract text from PDF. Please ensure the PDF contains readable text."
                )
        except ImportError:
            raise HTTPException(
                status_code=500,
                detail="PDF parsing library not installed. Please install PyPDF2."
            )
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise HTTPException(
                status_code=500,
                detail=f"Failed to extract text from PDF: {str(e)}"
            )

        # Store in Supabase
        result = supabase.client.table("knowledge_base_documents").insert({
            "filename": file.filename,
            "file_size": file_size,
            "content": text_content,
            "description": description,
            "uploaded_by_user_id": current_user["id"],
            "is_active": True
        }).execute()

        if result.data:
            return KnowledgeBaseUploadResponse(
                success=True,
                document_id=result.data[0]["id"],
                message=f"Document '{file.filename}' uploaded successfully"
            )
        else:
            raise HTTPException(status_code=500, detail="Failed to save document")

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading knowledge base document: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/knowledge-base/documents")
def get_knowledge_base_documents(
    current_user: Dict[str, Any] = Depends(get_current_user)
) -> List[KnowledgeBaseDocument]:
    """Get all active knowledge base documents."""
    if not supabase.is_connected():
        raise HTTPException(
            status_code=503,
            detail="Knowledge base requires Supabase connection"
        )

    try:
        result = supabase.client.table("knowledge_base_documents")\
            .select("id, filename, file_size, description, uploaded_by_user_id, is_active, created_at, updated_at")\
            .eq("is_active", True)\
            .order("created_at", desc=True)\
            .execute()

        return [KnowledgeBaseDocument(**doc) for doc in result.data]

    except Exception as e:
        logger.error(f"Error fetching knowledge base documents: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/knowledge-base/documents/{document_id}")
def delete_knowledge_base_document(
    document_id: str,
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """
    Delete (deactivate) a knowledge base document.
    Only the uploader can delete their documents.
    """
    if not supabase.is_connected():
        raise HTTPException(
            status_code=503,
            detail="Knowledge base requires Supabase connection"
        )

    try:
        # Soft delete by setting is_active to false
        result = supabase.client.table("knowledge_base_documents")\
            .update({"is_active": False})\
            .eq("id", document_id)\
            .eq("uploaded_by_user_id", current_user["id"])\
            .execute()

        if result.data:
            return {"success": True, "message": "Document deleted"}
        else:
            raise HTTPException(status_code=404, detail="Document not found or access denied")

    except Exception as e:
        logger.error(f"Error deleting document {document_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# MICROSOFT TEAMS INTEGRATION ENDPOINTS
# ============================================================================

@app.get("/teams/members")
async def get_teams_members(
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """
    Get list of all users in the organization (Azure AD directory).
    Requires user to have authorized with User.Read.All permission.
    """
    try:
        # Use user ID as identifier (same as how tokens are stored)
        user_identifier = str(current_user.get("id"))

        # Get user's Outlook token (includes Teams permissions)
        outlook = get_outlook_client()
        tokens = outlook.get_user_auth_tokens(user_identifier)

        if not tokens:
            raise HTTPException(
                status_code=401,
                detail="Email not authorized. Please connect your email first."
            )

        access_token = tokens.get("access_token")

        # Fetch all organization users
        members = outlook.get_organization_users(access_token)

        return {
            "success": True,
            "members": members,
            "count": len(members)
        }

    except RuntimeError as e:
        error_msg = str(e)
        if "insufficient permissions" in error_msg.lower():
            raise HTTPException(
                status_code=403,
                detail="Please re-authorize your email connection to grant Teams permissions"
            )
        raise HTTPException(status_code=500, detail=error_msg)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching organization users: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/teams/send-assignment-notification")
async def send_teams_assignment_notification(
    assignee_user_id: str = Body(...),
    assignee_name: str = Body(...),
    lead_subject: str = Body(...),
    lead_email: str = Body(...),
    lead_company: Optional[str] = Body(None),
    notes: Optional[str] = Body(None),
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """
    Send a Teams notification to a user about a lead assignment.
    Requires user to have authorized with Teams permissions.
    """
    try:
        # Use user ID as identifier (same as how tokens are stored)
        user_identifier = str(current_user.get("id"))

        # Get user's Outlook token (includes Teams permissions)
        outlook = get_outlook_client()
        tokens = outlook.get_user_auth_tokens(user_identifier)

        if not tokens:
            raise HTTPException(
                status_code=401,
                detail="Email not authorized. Please connect your email first."
            )

        access_token = tokens.get("access_token")
        cfg = Config()
        app_url = cfg.FRONTEND_URL

        # Send Teams notification
        success = outlook.send_lead_assignment_notification(
            access_token=access_token,
            assignee_user_id=assignee_user_id,
            assignee_name=assignee_name,
            lead_subject=lead_subject,
            lead_email=lead_email,
            lead_company=lead_company,
            notes=notes,
            app_url=app_url
        )

        if not success:
            raise HTTPException(
                status_code=500,
                detail="Failed to send Teams notification"
            )

        return {
            "success": True,
            "message": f"Teams notification sent to {assignee_name}"
        }

    except RuntimeError as e:
        error_msg = str(e)
        if "insufficient permissions" in error_msg.lower():
            raise HTTPException(
                status_code=403,
                detail="Please re-authorize your email connection to grant Teams permissions"
            )
        raise HTTPException(status_code=500, detail=error_msg)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error sending Teams notification: {e}")
        raise HTTPException(status_code=500, detail=str(e))
